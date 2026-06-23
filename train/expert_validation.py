# -*- coding: utf-8 -*-
"""Tạo PHIẾU THẨM ĐỊNH reranker: chạy pipeline tốt nhất (bge-m3 + từ điển + ngữ nghĩa +
reranker fine-tune) trên các câu hỏi đời thường, ghi lại top-1 reranker CHỌN + luật gốc
gán nhãn, để chuyên gia pháp lý đánh giá lựa chọn nào phục vụ người dùng tốt hơn.
Trực tiếp kiểm chứng cảnh báo trung thực ở Mục 4.17. Xuất report/PhieuThamDinh_Reranker.docx.
Chạy: python train/expert_validation.py
"""
import sys, os, json, pickle
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.stdout.reconfigure(encoding="utf-8")
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
golden = json.load(open(os.path.join(ROOT, "eval", "golden_questions.json"), encoding="utf-8"))
NAT = [i for i, q in enumerate(golden) if q.get("type") == "natural"]
WHERE = {"effect_status": "In effect"}
FT = os.path.join(ROOT, "models", "ft-reranker-bge")
OUT_JSON = os.path.join(os.path.dirname(os.path.abspath(__file__)), "expert_validation.json")
OUT_DOCX = os.path.join(ROOT, "report", "PhieuThamDinh_Reranker.docx")


def gen_data():
    from rag_core import Retriever
    meta = pickle.load(open(os.path.join(ROOT, "bm25_meta.pkl"), "rb"))
    dn2t = {}
    for m in meta:
        dn = m.get("document_number")
        if dn and dn not in dn2t:
            dn2t[dn] = m.get("title", "") or ""
    r = Retriever(collection="bge_m3", model="BAAI/bge-m3", segment=False, expand=True, fp16=True,
                  sem_expand=True, sem_K=3, sem_tau=0.45, use_rerank=True, rerank_model=FT, n_candidates=30)
    rows = []
    for n, i in enumerate(NAT):
        q = golden[i]["q"]; rel = golden[i]["relevant"]
        hits = r.search(q, k=3, where=WHERE)
        top = [(h["metadata"].get("document_number"), h["metadata"].get("title", "")) for h in hits]
        rel_titles = [(dn, dn2t.get(dn, "")) for dn in rel]
        rows.append({"q": q, "top": top, "rel": rel_titles,
                     "hit": bool(top) and top[0][0] in set(rel)})
        if (n + 1) % 40 == 0:
            print(f"  {n+1}/{len(NAT)}", flush=True)
    json.dump(rows, open(OUT_JSON, "w", encoding="utf-8"), ensure_ascii=False, indent=1)
    print(f"Đã lưu {OUT_JSON} ({len(rows)} câu)")
    return rows


def make_docx(rows):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    s = doc.sections[0]
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = s.page_height, s.page_width
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(s, m, Cm(1.5))

    def style_run(rn, sz=11, bold=False, italic=False, color=None):
        rn.font.name = "Times New Roman"; rn.font.size = Pt(sz)
        rn.font.bold = bold; rn.font.italic = italic
        if color: rn.font.color.rgb = color

    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(h.add_run("PHIẾU THẨM ĐỊNH KẾT QUẢ RERANKER (CHUYÊN GIA PHÁP LÝ)"), 14, bold=True)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(sub.add_run("Đề tài: Hệ thống RAG tra cứu văn bản pháp luật/y tế, Cấu hình bge-m3 + từ điển + ngữ nghĩa + reranker fine-tune"), 11, italic=True)
    intro = doc.add_paragraph()
    style_run(intro.add_run("Hướng dẫn: với mỗi câu hỏi đời thường, cột “Văn bản hệ thống xếp #1” là văn bản reranker đưa lên đầu; cột “Luật gốc (nhãn)” là luật điều chỉnh được gán nhãn. Xin chuyên gia đánh giá ở cột cuối: (A) #1 phù hợp nhất; (B) Luật gốc phù hợp hơn; (C) cả hai đều dùng được; (D) cả hai chưa đúng, và ghi chú nếu cần. Mục tiêu: kiểm chứng việc reranker ưu tiên văn bản cấp Luật có thực sự phục vụ người dùng tốt hơn không."), 10.5)

    n_hit = sum(1 for r in rows if r["hit"])
    stat = doc.add_paragraph()
    style_run(stat.add_run(f"Tổng số câu: {len(rows)}. Số câu #1 trùng luật gốc gán nhãn: {n_hit} ({n_hit*100//len(rows)}%). "
                           f"Các câu còn lại #1 thường là nghị định/thông tư hướng dẫn, cần chuyên gia xác nhận mức phù hợp."), 10.5, italic=True)

    cols = ["STT", "Câu hỏi (đời thường)", "Văn bản hệ thống xếp #1", "Luật gốc (nhãn)", "Đánh giá (A/B/C/D)", "Ghi chú"]
    widths = [1.0, 6.5, 8.0, 6.5, 3.0, 4.0]
    t = doc.add_table(rows=1, cols=len(cols)); t.style = "Table Grid"
    for c, (name, w) in enumerate(zip(cols, widths)):
        cell = t.rows[0].cells[c]; cell.width = Cm(w)
        p = cell.paragraphs[0]; style_run(p.add_run(name), 10.5, bold=True)
    for idx, r in enumerate(rows, 1):
        cells = t.add_row().cells
        top1 = r["top"][0] if r["top"] else ("", "")
        rel = r["rel"][0] if r["rel"] else ("", "")
        vals = [str(idx), r["q"],
                f"{top1[0]}, {top1[1][:90]}",
                f"{rel[0]}, {rel[1][:80]}" + (f" (+{len(r['rel'])-1} VB)" if len(r["rel"]) > 1 else ""),
                "", ""]
        for c, (v, w) in enumerate(zip(vals, widths)):
            cells[c].width = Cm(w)
            p = cells[c].paragraphs[0]
            run = p.add_run(v)
            mark = (top1[0] in {d for d, _ in r["rel"]})
            style_run(run, 9.5, color=RGBColor(0x0b, 0x6b, 0x2e) if (c == 2 and mark) else None)
    end = doc.add_paragraph(); end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    end.paragraph_format.space_before = Pt(18)
    style_run(end.add_run("Người thẩm định (ký, ghi rõ họ tên): ……………………………………"), 11)
    doc.save(OUT_DOCX)
    print(f"✅ Đã lưu {OUT_DOCX}")


def main():
    if os.path.exists(OUT_JSON):
        rows = json.load(open(OUT_JSON, encoding="utf-8"))
        print(f"Nạp lại {OUT_JSON} ({len(rows)} câu), bỏ qua chạy pipeline.")
    else:
        rows = gen_data()
    make_docx(rows)


if __name__ == "__main__":
    main()
