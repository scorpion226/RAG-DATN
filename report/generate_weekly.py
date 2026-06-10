# -*- coding: utf-8 -*-
"""Sinh Báo cáo tiến độ hàng tuần (.docx) — thống kê nhiệm vụ & công việc đã thực hiện."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "BaoCao_TienDo_HangTuan.docx")
FONT = "Times New Roman"


def sf(run, size=13, bold=False, italic=False, color=None):
    run.font.name = FONT; run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), FONT)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); tcPr.append(sh)


def main():
    doc = Document()
    n = doc.styles['Normal']; n.font.name = FONT; n.font.size = Pt(13)
    n.element.rPr.rFonts.set(qn('w:cs'), FONT)
    s = doc.sections[0]
    s.page_width = Cm(21); s.page_height = Cm(29.7)
    s.top_margin = Cm(2); s.bottom_margin = Cm(2); s.left_margin = Cm(3); s.right_margin = Cm(2)

    def center(t, size, bold=True, after=4, before=0):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
        sf(p.add_run(t), size, bold=bold)

    def H(t):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
        sf(p.add_run(t), 13, bold=True); return p

    def body(t, italic=False):
        p = doc.add_paragraph(); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.space_after = Pt(4)
        sf(p.add_run(t), 13, italic=italic); return p

    def info(label, val):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        sf(p.add_run(label + ": "), 13, bold=True); sf(p.add_run(val), 13)

    # ---- Tiêu đề ----
    center("TRƯỜNG ĐẠI HỌC THỦY LỢI — KHOA CÔNG NGHỆ THÔNG TIN", 12, True, 2)
    center("BÁO CÁO TIẾN ĐỘ THỰC HIỆN ĐỒ ÁN TỐT NGHIỆP", 15, True, 12, 6)
    center("(Thống kê nhiệm vụ và công việc đã thực hiện)", 12, False, 14)
    info("Sinh viên", "Nguyễn Minh Hiếu")
    info("Ngành", "Công nghệ thông tin")
    info("Đề tài", "Hệ thống RAG hỗ trợ tra cứu văn bản pháp luật ngành y tế Việt Nam")
    info("Giảng viên hướng dẫn", "..............................")

    # ---- I. THỐNG KÊ TỔNG QUAN ----
    H("I. THỐNG KÊ TỔNG QUAN")
    stats = [
        ["Thời gian thực hiện", "12 tuần"],
        ["Quy mô dữ liệu", "21.490 văn bản y tế còn hiệu lực → 367.462 đoạn (chunk) đã làm sạch"],
        ["Mã nguồn", "15 module Python (~1.715 dòng) + giao diện web React"],
        ["Mô hình tích hợp", "PhoBERT bi-encoder (embedding), PhoRanker (reranker), PhoGPT-4B-Chat (sinh)"],
        ["Số thí nghiệm đánh giá", "04 thí nghiệm A/B (tách từ, reranker, kích thước chunk, hybrid)"],
        ["Kết quả truy xuất tốt nhất", "MRR = 0,722; Hit@10 = 1,000 (cấu hình Hybrid + Reranker)"],
        ["Sản phẩm", "Pipeline RAG hoàn chỉnh + chatbot web (chạy offline) + báo cáo"],
        ["Tiến độ tổng thể", "Hoàn thành hệ thống đầu–cuối; đang hoàn thiện báo cáo"],
    ]
    t = doc.add_table(rows=0, cols=2); t.style = 'Table Grid'
    for i, (a, b) in enumerate(stats):
        cells = t.add_row().cells
        cells[0].width = Cm(5.5); cells[1].width = Cm(10.5)
        sf(cells[0].paragraphs[0].add_run(a), 12, bold=True)
        sf(cells[1].paragraphs[0].add_run(b), 12)
        if i == 0:
            shade(cells[0], "EAF1FB"); shade(cells[1], "EAF1FB")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---- II. BÁO CÁO TIẾN ĐỘ THEO TUẦN ----
    H("II. CHI TIẾT TIẾN ĐỘ THEO TUẦN")
    weeks = [
        ("1", "Khảo sát, phân tích bài toán; xác định đề tài và phạm vi.",
         "Chốt đề tài RAG tra cứu pháp luật y tế; xác định 3 câu hỏi nghiên cứu.", "Hoàn thành"),
        ("2–3", "Nghiên cứu lý thuyết RAG, embedding, vector database; thu thập dữ liệu.",
         "Nắm kiến trúc RAG; chọn nguồn dữ liệu pháp luật công khai (Hugging Face).", "Hoàn thành"),
        ("4", "Tiền xử lý dữ liệu; lựa chọn mô hình.",
         "Chốt stack: PhoBERT (embedding), ChromaDB (vector DB), PhoGPT (LLM).", "Hoàn thành"),
        ("5–7", "Xây module lập chỉ mục; pipeline RAG cơ bản.",
         "Pipeline truy vấn được câu trả lời trên bộ dữ liệu nhỏ.", "Hoàn thành"),
        ("8", "Kiểm định & khắc phục lỗi dữ liệu; xây lại pipeline; embedding toàn corpus.",
         "Phát hiện lỗi nhân bản 27,7 triệu dòng (98,7% trùng) và mất 75% văn bản; rebuild sạch 367.462 đoạn; embedding đầy đủ vào ChromaDB.", "Hoàn thành"),
        ("9", "Tích hợp LLM PhoGPT; xây backend và giao diện.",
         "Chạy PhoGPT-4B-Chat lượng tử hóa (GGUF Q4, llama.cpp) trên CPU; backend FastAPI + web chat sinh câu trả lời có trích dẫn.", "Hoàn thành"),
        ("10", "Nâng cấp truy xuất: reranker và hybrid search.",
         "Thêm reranker PhoRanker và hybrid BM25+vector (RRF) kèm nhận diện số hiệu văn bản; cải thiện rõ độ chính xác.", "Hoàn thành"),
        ("11", "Đánh giá định lượng; nâng cấp giao diện.",
         "Xây bộ 18 câu hỏi vàng; đo Hit@k/MRR; 4 thí nghiệm A/B; giao diện React (lịch sử hội thoại, chạy offline).", "Hoàn thành"),
        ("12", "Viết báo cáo theo chuẩn trình bày của trường.",
         "Soạn báo cáo đầy đủ (Chương 0–5, tài liệu tham khảo IEEE, hình/bảng); chuẩn định dạng ĐH Thủy lợi.", "Đang thực hiện"),
    ]
    tb = doc.add_table(rows=1, cols=4); tb.style = 'Table Grid'
    heads = ["Tuần", "Nhiệm vụ", "Công việc / kết quả đã thực hiện", "Trạng thái"]
    wcm = [1.4, 4.3, 8.3, 2.0]
    for i, h in enumerate(heads):
        c = tb.rows[0].cells[i]; c.width = Cm(wcm[i]); shade(c, "D9E2F3")
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sf(c.paragraphs[0].add_run(h), 12, bold=True)
    for wk, nv, kq, st in weeks:
        cells = tb.add_row().cells
        for i, v in enumerate([wk, nv, kq, st]):
            cells[i].width = Cm(wcm[i])
            if i in (0, 3): cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            sf(cells[i].paragraphs[0].add_run(v), 12)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---- III. KHÓ KHĂN & HƯỚNG XỬ LÝ ----
    H("III. KHÓ KHĂN GẶP PHẢI VÀ HƯỚNG XỬ LÝ")
    diff = [
        ("Dữ liệu bị nhân bản và mất mát do pipeline cũ.",
         "Đo tỷ lệ trùng, truy nguyên lỗi vòng lặp ghi, xây lại pipeline sạch (chunk theo Điều, khử trùng lặp)."),
        ("Máy chỉ có CPU 16GB, không chạy nổi LLM 4B ở dạng đầy đủ.",
         "Dùng PhoGPT bản lượng tử hóa GGUF Q4 (2,36GB) qua llama.cpp; xử lý dữ liệu theo luồng/batch."),
        ("Embedding toàn corpus rất lâu (~13 giờ trên CPU).",
         "Cho phép lập chỉ mục resume; chạy nền qua đêm, không phụ thuộc một phiên liên tục."),
        ("Vector search yếu khi hỏi theo số hiệu văn bản.",
         "Bổ sung BM25 (hybrid, RRF) và nhận diện số hiệu để truy vấn đúng văn bản."),
        ("Máy chưa cài Node.js để dựng React.",
         "Dùng React qua thư viện đóng gói cục bộ (không cần build), giao diện chạy offline."),
    ]
    td = doc.add_table(rows=1, cols=2); td.style = 'Table Grid'
    for i, h in enumerate(["Khó khăn", "Hướng xử lý"]):
        c = td.rows[0].cells[i]; c.width = Cm(8); shade(c, "D9E2F3")
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sf(c.paragraphs[0].add_run(h), 12, bold=True)
    for a, b in diff:
        cells = td.add_row().cells
        cells[0].width = Cm(8); cells[1].width = Cm(8)
        sf(cells[0].paragraphs[0].add_run(a), 12); sf(cells[1].paragraphs[0].add_run(b), 12)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---- IV. KẾ HOẠCH TIẾP THEO ----
    H("IV. KẾ HOẠCH TUẦN TIẾP THEO")
    for t_ in ["Hoàn thiện báo cáo và chỉnh sửa theo góp ý của giảng viên hướng dẫn.",
               "Mở rộng bộ câu hỏi đánh giá và bổ sung đo độ trung thực (faithfulness) của câu trả lời.",
               "Chuẩn bị slide và kịch bản trình diễn (demo) hệ thống cho buổi bảo vệ."]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        sf(p.add_run(t_), 13)

    # xác nhận
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sf(p.add_run("Hà Nội, ngày ... tháng ... năm 2026"), 13, italic=True)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sf(p2.add_run("Sinh viên thực hiện"), 13, bold=True)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.paragraph_format.space_before = Pt(36)
    sf(p3.add_run("Nguyễn Minh Hiếu"), 13, bold=True)

    try:
        z = doc.settings.element.find(qn('w:zoom'))
        if z is not None and z.get(qn('w:percent')) is None: z.set(qn('w:percent'), '100')
    except Exception: pass

    doc.save(OUT)
    try: print("Da luu:", OUT)
    except Exception: pass


if __name__ == "__main__":
    main()
