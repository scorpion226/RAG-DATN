# -*- coding: utf-8 -*-
"""Báo cáo tiến độ hàng tuần (NGẮN GỌN) — bám thước đo bảng kế hoạch 14 tuần."""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "BaoCao_HangTuan_NganGon.docx")
FONT = "Times New Roman"


def sf(run, size=13, bold=False, italic=False):
    run.font.name = FONT; run.font.size = Pt(size); run.bold = bold; run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), FONT)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); tcPr.append(sh)


def ctext(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    sf(p.add_run(text), size, bold=bold)


# Thời gian | Nội dung thực hiện | Kết quả đạt được | Trạng thái
ROWS = [
    ("Tuần 1\n23/03 - 29/03",
     "Khảo sát, phân tích bài toán; xác định phạm vi và nguồn dữ liệu.",
     "Chốt đề tài RAG; chọn nguồn dữ liệu pháp luật công khai (Hugging Face).",
     "Hoàn thành"),
    ("Tuần 2 - 3\n30/03 - 12/04",
     "Nghiên cứu RAG, embedding, LLM tiếng Việt; thu thập và tiền xử lý dữ liệu.",
     "Lọc 21.490 văn bản y tế còn hiệu lực; làm sạch, phân đoạn 367.462 chunk.",
     "Hoàn thành"),
    ("Tuần 4 - 5\n13/04 - 26/04",
     "Lựa chọn mô hình, cơ sở dữ liệu; xây module lập chỉ mục.",
     "Embedding PhoBERT + ChromaDB; chỉ mục vector hoạt động.",
     "Hoàn thành"),
    ("Tuần 6 - 7\n27/04 - 10/05",
     "Xây pipeline RAG: truy vấn ngữ nghĩa và sinh câu trả lời.",
     "Pipeline hoàn chỉnh; tích hợp PhoGPT, trả lời kèm trích dẫn nguồn.",
     "Hoàn thành"),
    ("Tuần 8 - 9\n11/05 - 24/05",
     "Phát triển giao diện web; tối ưu truy vấn/sinh; tích hợp trích dẫn.",
     "Web demo React; thêm reranker và hybrid BM25; hiển thị nguồn trích dẫn.",
     "Hoàn thành"),
    ("Tuần 10\n25/05 - 31/05",
     "Xây bộ câu hỏi kiểm thử; thiết kế kịch bản đánh giá.",
     "Bộ câu hỏi vàng có đáp án tham chiếu; kế hoạch đánh giá định lượng.",
     "Hoàn thành"),
    ("Tuần 11 - 12\n01/06 - 14/06",
     "Thực nghiệm đánh giá; phân tích chất lượng câu trả lời.",
     "4 thí nghiệm A/B; MRR 0,722, Hit@10 1,000; bảng số liệu và biểu đồ.",
     "Đang thực hiện"),
    ("Tuần 13\n15/06 - 21/06",
     "Tổng hợp kết quả, hoàn thiện demo; chỉnh sửa pipeline.",
     "Demo hoàn chỉnh, sẵn sàng trình diễn.",
     "Kế hoạch"),
    ("Tuần 14\n22/06 - 28/06",
     "Hoàn thiện báo cáo, chuẩn bị slide và bảo vệ.",
     "Báo cáo đầy đủ, slide, mã nguồn và tài liệu hướng dẫn.",
     "Kế hoạch"),
]
STATUS_FILL = {"Hoàn thành": "E2EFDA", "Đang thực hiện": "FFF2CC", "Kế hoạch": "F2F2F2"}


def main():
    doc = Document()
    n = doc.styles['Normal']; n.font.name = FONT; n.font.size = Pt(13)
    n.element.rPr.rFonts.set(qn('w:cs'), FONT)
    s = doc.sections[0]
    s.page_width = Cm(21); s.page_height = Cm(29.7)
    s.top_margin = Cm(2); s.bottom_margin = Cm(2); s.left_margin = Cm(3); s.right_margin = Cm(2)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(t.add_run("BÁO CÁO TIẾN ĐỘ HÀNG TUẦN"), 15, bold=True)
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER; s2.paragraph_format.space_after = Pt(2)
    sf(s2.add_run("Đề tài: Hệ thống RAG hỗ trợ tra cứu văn bản pháp luật / y tế Việt Nam"), 12, italic=True)
    s3 = doc.add_paragraph(); s3.alignment = WD_ALIGN_PARAGRAPH.CENTER; s3.paragraph_format.space_after = Pt(8)
    sf(s3.add_run("Sinh viên: Nguyễn Minh Hiếu     —     GVHD: ......................"), 12, italic=True)

    widths = [Cm(2.9), Cm(5.3), Cm(5.4), Cm(2.4)]
    table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Thời gian", "Nội dung thực hiện", "Kết quả đạt được", "Trạng thái"]):
        c = table.rows[0].cells[i]; c.width = widths[i]; shade(c, "D9E2F3")
        ctext(c, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for tg, nd, kq, st in ROWS:
        cells = table.add_row().cells
        ctext(cells[0], tg, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        ctext(cells[1], nd)
        ctext(cells[2], kq)
        ctext(cells[3], st, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(cells[3], STATUS_FILL.get(st, "FFFFFF"))
        for i, c in enumerate(cells): c.width = widths[i]
    for row in table.rows:
        for i, c in enumerate(row.cells): c.width = widths[i]

    # chú thích trạng thái
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    note = doc.add_paragraph()
    sf(note.add_run("Ghi chú: "), 11, bold=True)
    sf(note.add_run("Hoàn thành (xanh) · Đang thực hiện (vàng) · Kế hoạch (xám). Cập nhật đến ngày 06/06/2026."), 11, italic=True)

    try:
        z = doc.settings.element.find(qn('w:zoom'))
        if z is not None and z.get(qn('w:percent')) is None: z.set(qn('w:percent'), '100')
    except Exception: pass

    doc.save(OUT)
    try: print("Da luu:", OUT)
    except Exception: pass


if __name__ == "__main__":
    main()
