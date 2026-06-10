# -*- coding: utf-8 -*-
"""Báo cáo công việc theo tuần (chi tiết, theo mẫu) cho đồ án RAG pháp luật/y tế."""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "BaoCao_CongViec_TheoTuan.docx")
FONT = "Times New Roman"


def sf(run, size=13, bold=False, italic=False):
    run.font.name = FONT; run.font.size = Pt(size); run.bold = bold; run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), FONT)


def main():
    doc = Document()
    n = doc.styles['Normal']; n.font.name = FONT; n.font.size = Pt(13)
    n.element.rPr.rFonts.set(qn('w:cs'), FONT)
    s = doc.sections[0]
    s.page_width = Cm(21); s.page_height = Cm(29.7)
    s.top_margin = Cm(2); s.bottom_margin = Cm(2); s.left_margin = Cm(3); s.right_margin = Cm(2)

    def para(text, italic=False, indent=0.0, after=4, justify=True):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE; pf.space_after = Pt(after)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
        if indent: pf.left_indent = Cm(indent)
        sf(p.add_run(text), 13, italic=italic); return p

    def lead(label, text, indent=0.5):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE; pf.space_after = Pt(3)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; pf.left_indent = Cm(indent)
        sf(p.add_run(label + " "), 13, bold=True); sf(p.add_run(text), 13)

    def li(text, indent=1.0):
        p = doc.add_paragraph(style='List Bullet'); pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE; pf.space_after = Pt(2)
        pf.left_indent = Cm(indent)
        sf(p.add_run(text), 13)

    def week(title):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.space_before = Pt(10); pf.space_after = Pt(3)
        sf(p.add_run(title), 13, bold=True)

    # ===== Tiêu đề =====
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(t.add_run("BÁO CÁO CÔNG VIỆC THEO TUẦN"), 15, bold=True)
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(s2.add_run("Đề tài: Hệ thống RAG hỗ trợ tra cứu văn bản pháp luật ngành y tế Việt Nam"), 12, italic=True)
    s3 = doc.add_paragraph(); s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(s3.add_run("Sinh viên: Nguyễn Minh Hiếu"), 12, italic=True)

    head = doc.add_paragraph(); head.paragraph_format.space_before = Pt(8)
    sf(head.add_run("Công việc đã thực hiện"), 13, bold=True)

    # ===== TUẦN 1 =====
    week("Tuần 1 (23/3 – 29/3): Khảo sát, phân tích bài toán và lựa chọn đề tài")
    lead("BG:", "Hệ thống văn bản pháp luật ngành y tế Việt Nam rất lớn và biến động liên tục (sửa đổi, hợp nhất, hết hiệu lực), ngôn ngữ pháp lý phức tạp với cấu trúc “Điều – Khoản – Điểm” đặc thù. Người dân, cán bộ y tế và sinh viên gặp khó khi tra cứu chính xác và nhanh chóng.")
    lead("Mov:", "Một hệ thống hỏi–đáp hiểu được ý định câu hỏi và trả lời kèm trích dẫn nguồn sẽ giúp tra cứu hiệu quả, tăng độ tin cậy. Đây là nhu cầu thực tế và có ý nghĩa ứng dụng cao.")
    lead("Challenge:", "Tìm kiếm từ khóa truyền thống chỉ khớp mặt chữ nên bỏ sót khi diễn đạt khác; còn hỏi trực tiếp một LLM thì dễ “bịa” và không truy được nguồn. Cần phương pháp vừa hiểu ngữ nghĩa vừa bám văn bản gốc.")
    para("Để giải quyết, đề tài sử dụng kiến trúc RAG (Retrieval-Augmented Generation): truy xuất đoạn văn bản liên quan từ kho tri thức thật rồi để mô hình ngôn ngữ sinh câu trả lời dựa trên đó và trích dẫn nguồn. So với fine-tuning toàn bộ tri thức vào mô hình, RAG cập nhật dữ liệu dễ, có trích dẫn và giảm bịa đặt.")
    para("Các hướng khác không được chọn vì: chỉ dùng LLM thuần dễ bịa đặt và không nguồn; chỉ dùng tìm kiếm từ khóa thì không hiểu ngữ nghĩa; fine-tuning một LLM lớn cho miền pháp luật tốn tài nguyên và khó cập nhật. Vì vậy em chọn đề tài “Hệ thống RAG hỗ trợ tra cứu văn bản pháp luật ngành y tế Việt Nam”.")
    lead("Input:", "câu hỏi tiếng Việt về pháp luật ngành y tế.")
    lead("Output:", "câu trả lời súc tích, bám văn bản, kèm trích dẫn số hiệu văn bản nguồn.")

    # ===== TUẦN 2 =====
    week("Tuần 2 (30/3 – 5/4): Nghiên cứu tổng quan về RAG và thu thập dữ liệu")
    para("Nghiên cứu kiến trúc RAG và các thành phần: embedding (mã hóa văn bản thành vector ngữ nghĩa), cơ sở dữ liệu vector (tìm kiếm lân cận), mô hình ngôn ngữ sinh. Xác định stack phù hợp tài nguyên CPU 16GB: PhoBERT (embedding tiếng Việt), ChromaDB (vector DB), PhoGPT (LLM tiếng Việt).")
    lead("Thu thập dataset:", "vietnamese-legal-documents (Hugging Face).")
    li("Khoảng 518.601 bản ghi metadata và 518.235 bản ghi nội dung văn bản pháp luật Việt Nam.")
    li("Mỗi văn bản có: số hiệu, tiêu đề, loại văn bản, ngành/lĩnh vực, cơ quan ban hành, ngày ban hành, ngày hiệu lực, tình trạng hiệu lực.")
    li("Phù hợp để lọc ra tập văn bản ngành y tế phục vụ đề tài.")

    # ===== TUẦN 3 =====
    week("Tuần 3 (6/4 – 12/4): Tiền xử lý dữ liệu — lọc, làm sạch, khảo sát (EDA)")
    lead("EDA:", "Lọc các văn bản có ngành/lĩnh vực là “Y tế” thu được 34.223 văn bản, trong đó 34.166 văn bản có nội dung. Tiếp tục giữ văn bản còn hiệu lực được 21.490 văn bản. Nội dung là văn bản thuần, còn dấu xuống dòng và mốc “Điều”, “Chương”.")
    lead("Preprocessing:", "")
    li("Lọc theo ngành y tế và theo tình trạng còn hiệu lực (chatbot pháp luật chỉ nên tư vấn văn bản còn hiệu lực).")
    li("Làm sạch giữ lại cấu trúc xuống dòng và các dấu “/ - ( ) %” để bảo toàn số hiệu văn bản (vd 147/2025/NĐ-CP) — định danh để trích dẫn.")
    li("Kiểm định dữ liệu: phát hiện pipeline ban đầu bị lỗi nhân bản (27,7 triệu dòng, 98,7% trùng) và mất ~75% văn bản => xây lại pipeline sạch.")

    # ===== TUẦN 4 =====
    week("Tuần 4 (13/4 – 19/4): Chia đoạn (chunking) và xây tập chỉ mục")
    para("Chia văn bản thành đoạn (chunk) bằng thuật toán đệ quy, ưu tiên cắt tại mốc pháp lý “Điều” → “Chương” → “Mục” → câu, kích thước ~1200 ký tự, chồng lấp 150 ký tự. Cắt theo “Điều” giúp mỗi đoạn là đơn vị ngữ nghĩa hoàn chỉnh, thay vì cắt cứng làm đứt câu.")
    li("Khử trùng lặp đoạn (loại 9.131 đoạn boilerplate trùng).")
    li("Kết quả: 367.462 đoạn từ 21.490 văn bản, trung bình 17,1 đoạn/văn bản, ~1.036 ký tự/đoạn.")
    para("=> Có tập dữ liệu sạch sẵn sàng để mã hóa và lập chỉ mục.")

    # ===== TUẦN 5 =====
    week("Tuần 5 (20/4 – 26/4): Xây module embedding + ChromaDB, kiểm tra ban đầu")
    para("Dùng bi-encoder dựa trên PhoBERT (vietnamese-bi-encoder, vector 768 chiều) để mã hóa đoạn; tách từ tiếng Việt bằng pyvi trước khi mã hóa (vì PhoBERT huấn luyện trên dữ liệu đã tách từ). Lưu vector và metadata vào ChromaDB (độ đo cosine, chỉ mục HNSW).")
    li("Lập chỉ mục thử trên một tập nhỏ (15.000 đoạn) để kiểm tra luồng truy vấn.")
    li("Truy vấn thử cho kết quả đúng ngữ nghĩa (vd hỏi điều kiện hành nghề → trả về Điều 19 Luật Khám bệnh, chữa bệnh 2023).")
    para("=> Pipeline truy xuất cơ bản hoạt động trên dữ liệu nhỏ.")

    # ===== TUẦN 6 =====
    week("Tuần 6 (27/4 – 3/5): Lập chỉ mục toàn corpus + backend, giao diện cơ bản")
    li("Embedding toàn bộ 367.462 đoạn vào ChromaDB; tốc độ ~7,7 đoạn/giây trên CPU (~13 giờ), chạy nền qua đêm.")
    li("Cho phép lập chỉ mục resume: chạy lại bỏ qua phần đã nạp, không phụ thuộc một phiên liên tục.")
    li("Dựng backend FastAPI (điểm cuối /chat, /health) và giao diện web hỏi–đáp cơ bản.")

    # ===== TUẦN 7 =====
    week("Tuần 7 (4/5 – 10/5): Tích hợp mô hình sinh PhoGPT, có câu trả lời có trích dẫn")
    para("Tích hợp PhoGPT-4B-Chat ở dạng lượng tử hóa GGUF Q4_K_M (2,36GB) qua llama.cpp để chạy trên CPU. Prompt yêu cầu mô hình chỉ dựa vào ngữ cảnh, bắt buộc trích dẫn số hiệu và nói rõ khi không tìm thấy — nhằm giảm bịa đặt.")
    li("Thời gian sinh ~20–40 giây/câu trên CPU.")
    li("Câu trả lời bám ngữ cảnh, có trích dẫn số hiệu văn bản nguồn.")

    # ===== TUẦN 8 =====
    week("Tuần 8 (11/5 – 17/5): Nâng cấp truy xuất — Reranker (PhoRanker)")
    para("Thêm bước xếp hạng lại bằng cross-encoder PhoRanker: bi-encoder lấy nhanh 30 ứng viên, cross-encoder đọc đồng thời (câu hỏi, đoạn) để chấm chính xác hơn rồi chọn k tốt nhất.")
    li("Kết quả: Hit@1 tăng từ 0,389 lên 0,500; MRR tăng 0,058.")

    # ===== TUẦN 9 =====
    week("Tuần 9 (18/5 – 24/5): Hybrid search BM25 + vector và nhận diện số hiệu")
    para("Bổ sung tìm kiếm từ vựng BM25 và hợp nhất với vector bằng Reciprocal Rank Fusion (RRF), kèm nhận diện số hiệu văn bản trong câu hỏi để truy vấn đúng văn bản. Mục tiêu: bù điểm yếu của vector khi hỏi theo số hiệu (vd 96/2023/NĐ-CP).")
    li("Hybrid đạt Hit@10 = 1,000; kết hợp reranker đạt MRR = 0,722 (tốt nhất).")
    li("Truy vấn theo số hiệu trả về đúng văn bản, trong khi vector thuần thường trượt.")

    # ===== TUẦN 10 =====
    week("Tuần 10 (25/5 – 31/5): Đánh giá định lượng và các thí nghiệm A/B")
    para("Xây bộ 18 câu hỏi vàng (neo vào các luật y tế, có bổ sung văn bản hợp nhất tương đương) và đo Hit@k, MRR. Thực hiện 4 thí nghiệm A/B chứng minh đóng góp từng kỹ thuật.")
    li("Tách từ pyvi: Hit@3 tăng 0,556 → 0,778.")
    li("Reranker: Hit@1 tăng 0,389 → 0,500.")
    li("Kích thước chunk: cắt 500 ký tự tốn gấp ~2,3 lần số đoạn mà kém hơn → chọn cắt theo “Điều”.")
    li("Hybrid + Reranker: MRR 0,722; Hit@10 0,944–1,000.")

    # ===== TUẦN 11 =====
    week("Tuần 11 (1/6 – 7/6): Hoàn thiện giao diện và viết báo cáo (đang thực hiện)")
    li("Nâng cấp giao diện web bằng React: lịch sử hội thoại, thẻ nguồn, chạy offline (không cần Internet khi demo).")
    li("Soạn báo cáo theo chuẩn trình bày của trường (Chương 0–5, hình/bảng, tài liệu tham khảo IEEE).")
    para("=> Hệ thống đã hoàn chỉnh đầu–cuối; đang hoàn thiện báo cáo và chuẩn bị bảo vệ.")

    try:
        z = doc.settings.element.find(qn('w:zoom'))
        if z is not None and z.get(qn('w:percent')) is None: z.set(qn('w:percent'), '100')
    except Exception: pass

    doc.save(OUT)
    try: print("Da luu:", OUT)
    except Exception: pass


if __name__ == "__main__":
    main()
