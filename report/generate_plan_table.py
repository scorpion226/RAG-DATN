# -*- coding: utf-8 -*-
"""Bảng kế hoạch/tiến độ thực hiện đồ án (3 cột) — theo đúng mẫu yêu cầu."""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "KeHoach_TienDo_DoAn.docx")
FONT = "Times New Roman"


def sf(run, size=13, bold=False, italic=False):
    run.font.name = FONT; run.font.size = Pt(size); run.bold = bold; run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), FONT)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); tcPr.append(sh)


def cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    sf(p.add_run(text), 13, bold=bold)


ROWS = [
    ("Tuần 1: 23/03 - 29/03",
     "Khảo sát, phân tích bài toán tra cứu văn bản pháp luật/y tế; xác định phạm vi, nguồn dữ liệu",
     "Báo cáo khảo sát, danh sách nguồn dữ liệu chính thống và yêu cầu hệ thống"),
    ("Tuần 2 - Tuần 3: 30/03 - 12/04",
     "Nghiên cứu lý thuyết RAG, mô hình embedding, LLM tiếng Việt; thu thập và tiền xử lý dữ liệu",
     "Bộ dữ liệu văn bản pháp luật và y tế đã làm sạch, chuẩn hóa, phân đoạn; báo cáo tổng quan"),
    ("Tuần 4 - Tuần 5: 13/04 - 26/04",
     "Lựa chọn mô hình và cơ sở dữ liệu; xây dựng module lập chỉ mục",
     "Module tạo và lưu trữ trên database hoạt động"),
    ("Tuần 6 - Tuần 7: 27/04 - 10/05",
     "Xây dựng pipeline RAG: module truy vấn ngữ nghĩa và module sinh câu trả lời",
     "Pipeline cơ bản hoàn chỉnh; có thể truy vấn và nhận câu trả lời từ bộ dữ liệu nhỏ"),
    ("Tuần 8 - Tuần 9: 11/05 - 24/05",
     "Phát triển giao diện web, tối ưu hiệu năng truy vấn và sinh; tích hợp trích dẫn nguồn",
     "Web demo với giao diện hỏi đáp, hiển thị câu trả lời kèm trích dẫn văn bản nguồn"),
    ("Tuần 10: 25/05 - 31/05",
     "Xây dựng bộ câu hỏi kiểm thử cho hai lĩnh vực; thiết kế kịch bản đánh giá",
     "Bộ 50-100 câu hỏi mẫu có đáp án tham chiếu; kế hoạch đánh giá định lượng và định tính"),
    ("Tuần 11 - Tuần 12: 01/06 - 14/06",
     "Thực nghiệm đánh giá: tính Precision, Recall; đánh giá chất lượng câu trả lời",
     "Bảng số liệu, biểu đồ phân tích; đánh giá độ chính xác, mức độ trích dẫn, thời gian phản hồi"),
    ("Tuần 13: 15/06 - 21/06",
     "Tổng hợp kết quả, hoàn thiện sản phẩm demo; chỉnh sửa pipeline",
     "Demo hoàn chỉnh, sẵn sàng trình diễn"),
    ("Tuần 14: 22/06 - 28/06",
     "Hoàn thiện báo cáo đồ án, chuẩn bị slide thuyết trình và bảo vệ",
     "Báo cáo đồ án đầy đủ, slide thuyết trình, mã nguồn và tài liệu hướng dẫn"),
]


def main():
    doc = Document()
    n = doc.styles['Normal']; n.font.name = FONT; n.font.size = Pt(13)
    n.element.rPr.rFonts.set(qn('w:cs'), FONT)
    s = doc.sections[0]
    s.page_width = Cm(21); s.page_height = Cm(29.7)
    s.top_margin = Cm(2); s.bottom_margin = Cm(2); s.left_margin = Cm(3); s.right_margin = Cm(2)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(t.add_run("KẾ HOẠCH THỰC HIỆN ĐỒ ÁN TỐT NGHIỆP"), 15, bold=True)
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2.paragraph_format.space_after = Pt(10)
    sf(s2.add_run("Đề tài: Hệ thống RAG hỗ trợ tra cứu văn bản pháp luật / y tế Việt Nam"), 12, italic=True)

    widths = [Cm(3.6), Cm(6.4), Cm(6.0)]
    table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Thời gian thực hiện", "Nội dung thực hiện", "Kết quả đạt được"]):
        c = table.rows[0].cells[i]; c.width = widths[i]; shade(c, "D9E2F3")
        cell_text(c, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for tg, nd, kq in ROWS:
        cells = table.add_row().cells
        for i, v in enumerate([tg, nd, kq]):
            cells[i].width = widths[i]
            cell_text(cells[i], v, bold=(i == 0),
                      align=WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.LEFT)

    # ép lại độ rộng cột (tránh Word tự co)
    for row in table.rows:
        for i, c in enumerate(row.cells):
            c.width = widths[i]

    try:
        z = doc.settings.element.find(qn('w:zoom'))
        if z is not None and z.get(qn('w:percent')) is None: z.set(qn('w:percent'), '100')
    except Exception: pass

    doc.save(OUT)
    try: print("Da luu:", OUT)
    except Exception: pass


if __name__ == "__main__":
    main()
