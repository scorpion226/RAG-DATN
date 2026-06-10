# -*- coding: utf-8 -*-
"""Báo cáo công việc theo tuần (văn xuôi) — bổ sung T1-6, phân tích từng bước T7-12,
đối chiếu bảng tiến độ. Chèn hình minh họa từ report/figs/."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__)); FIGS = os.path.join(HERE, "figs")
OUT = os.path.join(HERE, "BaoCao_CongViec_TheoTuan.docx")
FONT = "Times New Roman"


def sf(run, size=13, bold=False, italic=False, color=None):
    run.font.name = FONT; run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), FONT)


doc = Document()
n = doc.styles['Normal']; n.font.name = FONT; n.font.size = Pt(13)
n.element.rPr.rFonts.set(qn('w:cs'), FONT)
s = doc.sections[0]; s.page_width = Cm(21); s.page_height = Cm(29.7)
s.top_margin = Cm(2); s.bottom_margin = Cm(2); s.left_margin = Cm(3); s.right_margin = Cm(2)


def week(title):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    sf(p.add_run(title), 13, bold=True)


def b(text, lead=None):
    p = doc.add_paragraph(style='List Bullet'); pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE; pf.space_after = Pt(2); pf.left_indent = Cm(1.0)
    if lead: sf(p.add_run(lead), 13, bold=True)
    sf(p.add_run(text), 13)


def para(text, indent=0.5, italic=False, after=3):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE; pf.space_after = Pt(after)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; pf.left_indent = Cm(indent)
    sf(p.add_run(text), 13, italic=italic)


def note(text):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(4); pf.left_indent = Cm(0.5)
    sf(p.add_run("→ Đối chiếu bảng tiến độ: "), 12, bold=True, italic=True, color=RGBColor(0x1f, 0x4e, 0x79))
    sf(p.add_run(text), 12, italic=True, color=RGBColor(0x1f, 0x4e, 0x79))


def fig(path, caption, w=14):
    if not os.path.exists(path): return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(4)
    p.add_run().add_picture(path, width=Cm(w))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(c.add_run(caption), 11, italic=True); c.paragraph_format.space_after = Pt(6)


def imgph(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf = p.paragraph_format
    pf.space_after = Pt(6)
    sf(p.add_run(f"[ Chèn ảnh: {text} ]"), 11, italic=True, color=RGBColor(0x80, 0x80, 0x80))


# ===== Tiêu đề =====
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(t.add_run("BÁO CÁO CÔNG VIỆC THEO TUẦN"), 15, bold=True)
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER; st.paragraph_format.space_after = Pt(2)
sf(st.add_run("Đề tài: Hệ thống RAG hỗ trợ tra cứu văn bản pháp luật / y tế Việt Nam"), 12, italic=True)
st2 = doc.add_paragraph(); st2.alignment = WD_ALIGN_PARAGRAPH.CENTER; st2.paragraph_format.space_after = Pt(8)
sf(st2.add_run("Sinh viên: Nguyễn Minh Hiếu"), 12, italic=True)
h = doc.add_paragraph(); sf(h.add_run("Công việc đã thực hiện"), 14, bold=True)

# ===== TUẦN 1 =====
week("Tuần 1 (23/3 – 29/3):")
b("Đã khảo sát, phân tích bài toán; xác định phạm vi và nguồn dữ liệu.")
b("Hoàn thành báo cáo khảo sát, danh sách nguồn dữ liệu chính thống và yêu cầu hệ thống.")
para("Bổ sung: Bài toán được xác định là xây dựng chatbot RAG hỏi–đáp văn bản pháp luật ngành y tế, chạy trên máy cá nhân (CPU, 16GB RAM). Yêu cầu hệ thống: trả lời có trích dẫn số hiệu văn bản, ưu tiên văn bản còn hiệu lực, có giao diện web. Nguồn dữ liệu chọn dùng là bộ vietnamese-legal-documents (Hugging Face) gồm ~518.000 văn bản pháp luật Việt Nam — lý do: công khai, có đầy đủ metadata (số hiệu, ngành, tình trạng hiệu lực).", italic=True)
para("Sản phẩm: Báo cáo khảo sát.")
note("khớp “Tuần 1: 23/03–29/03 — Khảo sát, xác định phạm vi, nguồn dữ liệu”. Trạng thái: Hoàn thành.")

# ===== TUẦN 2 =====
week("Tuần 2 (30/3 – 5/4):")
b("Đã nghiên cứu lý thuyết RAG, mô hình embedding, LLM tiếng Việt; thu thập dữ liệu pháp luật.")
b("Hoàn thành báo cáo tổng quan về lý thuyết, mô hình, LLM, CSDL và lý do lựa chọn.")
para("Bổ sung: Nắm kiến trúc RAG gồm hai khối Retriever (truy xuất) và Generator (sinh). Xác định bộ công cụ: embedding PhoBERT (bi-encoder), cơ sở dữ liệu vector ChromaDB, mô hình sinh PhoGPT — đều tối ưu cho tiếng Việt và chạy được trên CPU. Tiến hành tải và khảo sát bộ dữ liệu (khoảng 3,5GB nội dung và metadata).", italic=True)
fig(os.path.join(FIGS, "real_sectors.png"), "Hình 1. Phân bố lĩnh vực của bộ dữ liệu (Thể thao – Y tế: 34.166 văn bản)", 14)
fig(os.path.join(FIGS, "real_byyear.png"), "Hình 2. Phân bố số văn bản pháp luật theo năm", 14)
para("Sản phẩm: Báo cáo tổng quan.")
note("khớp “Tuần 2–3: 30/03–12/04 — Nghiên cứu lý thuyết, thu thập và tiền xử lý dữ liệu”. Trạng thái: Hoàn thành.")

# ===== TUẦN 3 =====
week("Tuần 3 (6/4 – 11/4):")
b("Đã chỉnh sửa nội dung đề tài theo hướng pháp luật đối với y tế.")
b("Đã hoàn thành thu thập và tiền xử lý dữ liệu: bộ dữ liệu văn bản pháp luật y tế đã làm sạch, chuẩn hóa, phân đoạn.")
para("Bổ sung: Lọc theo ngành y tế thu được 34.166 văn bản có nội dung; giữ văn bản còn hiệu lực còn 21.490 văn bản. Sau khi làm sạch (giữ số hiệu) và phân đoạn theo “Điều”, khử trùng lặp, thu được 367.462 đoạn (trung bình ~17 đoạn/văn bản, ~1.036 ký tự/đoạn). Trong quá trình kiểm định đã phát hiện và khắc phục lỗi nhân bản dữ liệu nghiêm trọng của pipeline ban đầu (27,7 triệu dòng, 98,7% trùng).", italic=True)
fig(os.path.join(FIGS, "real_doctype.png"), "Hình 3. Phân bố loại văn bản trong bộ dữ liệu", 14)
fig(os.path.join(FIGS, "sample_data.png"), "Hình 4. Một số mẫu dữ liệu sau khi làm sạch và phân đoạn", 15)
note("khớp “Tuần 2–3: thu thập và tiền xử lý dữ liệu → bộ dữ liệu đã làm sạch, chuẩn hóa, phân đoạn”. Trạng thái: Hoàn thành.")

# ===== TUẦN 4 =====
week("Tuần 4 (12/4 – 19/4):")
b("Tìm hiểu và phân tích lý do chọn mô hình embedding, mô hình sinh và cơ sở dữ liệu vector cho đồ án.")
para("Bổ sung: ", italic=True)
b("Embedding: dùng bi-encoder dựa trên PhoBERT (768 chiều). Lý do: PhoBERT gốc là masked-LM nên vector [CLS] không tối ưu cho so khớp câu; bi-encoder được tinh chỉnh cho tìm kiếm ngữ nghĩa.", "Mô hình embedding — ")
b("ChromaDB (chỉ mục HNSW, độ đo cosine). Lý do: nhẹ, lưu trên đĩa, lọc được theo metadata (ngành, hiệu lực).", "CSDL vector — ")
b("PhoGPT-4B-Chat bản lượng tử hóa GGUF Q4 (2,36GB) qua llama.cpp. Lý do: bản đầy đủ ~16GB không vừa máy; lượng tử hóa cho phép chạy CPU với mất mát nhỏ.", "Mô hình sinh — ")
imgph("ảnh báo cáo tuần 4 của em (so sánh/lý do chọn mô hình)")
note("khớp “Tuần 4–5: 13/04–26/04 — Lựa chọn mô hình và cơ sở dữ liệu”. Trạng thái: Hoàn thành.")

# ===== TUẦN 5 =====
week("Tuần 5 (20/4 – 27/4):")
b("Xây dựng module lập chỉ mục (indexing).")
para("Bổ sung: Module thực hiện: tách từ bằng pyvi → mã hóa embedding PhoBERT → ghi vector kèm metadata vào ChromaDB. Hỗ trợ lập chỉ mục theo lô (batch) và khả năng resume khi gián đoạn. Đã chạy thử trên tập nhỏ 15.000 đoạn để kiểm tra luồng hoạt động.", italic=True)
para("Kết quả: module tạo và lưu trữ vector trên cơ sở dữ liệu hoạt động đúng.")
note("khớp “Tuần 4–5: xây dựng module lập chỉ mục → module tạo và lưu trữ trên database hoạt động”. Trạng thái: Hoàn thành.")

# ===== TUẦN 6 =====
week("Tuần 6 (28/4 – 3/5):")
b("Nghiên cứu và xây dựng pipeline RAG.")
para("Bổ sung: Thiết kế luồng xử lý một câu hỏi: tách từ và mã hóa câu hỏi → truy xuất top-k đoạn gần nhất từ ChromaDB → dựng prompt (ngữ cảnh + yêu cầu trích dẫn) → đưa vào mô hình sinh. Đồng thời tiến hành lập chỉ mục toàn bộ 367.462 đoạn (tốc độ ~7,7 đoạn/giây trên CPU, ~13 giờ, chạy nền qua đêm, có resume).", italic=True)
fig(os.path.join(FIGS, "architecture.png"), "Hình 5. Kiến trúc/pipeline hệ thống RAG", 10)
note("khớp “Tuần 6–7: 27/04–10/05 — Xây dựng pipeline RAG”. Trạng thái: Hoàn thành.")

# ===== TUẦN 7 =====
week("Tuần 7 (4/5 – 11/5): Pipeline cơ bản — phân tích, đánh giá")
b("Hoàn thành pipeline cơ bản: có thể truy vấn và nhận câu trả lời từ bộ dữ liệu nhỏ.")
para("Phân tích từng bước:", italic=True)
b("Câu hỏi được tách từ và mã hóa cùng mô hình lúc lập chỉ mục để bảo đảm nhất quán phân phối embedding.", "Bước truy vấn — ")
b("Với câu hỏi “Điều kiện để cá nhân được phép khám bệnh, chữa bệnh?”, hệ thống trả về đúng Điều 19 Luật Khám bệnh, chữa bệnh 2023 (15/2023/QH15) với điểm tương đồng 0,83.", "Kết quả — ")
b("Truy xuất theo ngữ nghĩa chính xác trên tập nhỏ; pipeline đầu–cuối thông suốt. Đây là cơ sở để mở rộng ra toàn corpus và nâng cấp ở các tuần sau.", "Đánh giá — ")
note("khớp “Tuần 6–7 → Pipeline cơ bản hoàn chỉnh, có thể truy vấn và nhận câu trả lời từ bộ dữ liệu nhỏ”. Trạng thái: Hoàn thành.")

# ===== TUẦN 8 =====
week("Tuần 8 (12/5 – 18/5): Tích hợp mô hình sinh PhoGPT và giao diện — phân tích, đánh giá")
b("Tích hợp PhoGPT-4B-Chat (GGUF Q4) sinh câu trả lời; dựng backend FastAPI và giao diện web.")
para("Phân tích từng bước:", italic=True)
b("Prompt buộc mô hình chỉ dựa vào ngữ cảnh, trích dẫn số hiệu và nói rõ khi không có thông tin → giảm bịa đặt.", "Thiết kế prompt — ")
b("PhoGPT trả lời bám đúng nội dung Điều 19 và dẫn số hiệu 15/2023/QH15; thời gian sinh ~20–40 giây/câu trên CPU.", "Kết quả — ")
b("Phần sinh hoạt động ổn định; câu trả lời có trích dẫn, kiểm chứng được — đáp ứng yêu cầu của miền pháp luật.", "Đánh giá — ")
fig(os.path.join(FIGS, "web_result.png"), "Hình 6. Giao diện web demo — kết quả truy vấn thật: câu trả lời của PhoGPT kèm 5 nguồn trích dẫn (số hiệu, loại văn bản, tình trạng hiệu lực, điểm liên quan)", 12)
note("khớp “Tuần 8–9: 11/05–24/05 — Phát triển giao diện web; tích hợp trích dẫn nguồn”. Trạng thái: Hoàn thành.")

# ===== TUẦN 9 =====
week("Tuần 9 (19/5 – 25/5): Nâng cấp truy xuất (Reranker + Hybrid) — phân tích, đánh giá")
b("Thêm reranker PhoRanker; bổ sung tìm kiếm từ vựng BM25 và hợp nhất hybrid (RRF); nhận diện số hiệu văn bản.")
para("Phân tích từng bước (so sánh A/B):", italic=True)
b("cross-encoder xếp lại 30 ứng viên đầu → Hit@1 tăng 0,389 → 0,500; MRR tăng 0,058.", "Reranker — ")
b("BM25 + vector hợp nhất bằng RRF → Hit@10 đạt 1,000; truy vấn theo số hiệu (vd 96/2023/NĐ-CP) trả về đúng văn bản, trong khi vector thuần thường trượt.", "Hybrid — ")
b("mỗi kỹ thuật bổ sung đều cải thiện rõ chất lượng truy xuất; hybrid khắc phục triệt để điểm yếu của vector với truy vấn theo số hiệu.", "Đánh giá — ")
note("khớp “Tuần 8–9: tối ưu hiệu năng truy vấn và sinh”. Trạng thái: Hoàn thành.")

# ===== TUẦN 10 =====
week("Tuần 10 (26/5 – 1/6): Bộ câu hỏi kiểm thử và kịch bản đánh giá — phân tích")
b("Xây dựng bộ câu hỏi vàng và thiết kế kịch bản đánh giá định lượng.")
para("Phân tích từng bước:", italic=True)
b("18 câu hỏi neo vào các luật y tế trọng yếu (Khám chữa bệnh, Dược, An toàn thực phẩm, BHYT, thuốc lá, rượu bia, hiến mô tạng, HIV, phòng bệnh).", "Bộ câu hỏi — ")
b("nhãn “văn bản liên quan” được bổ sung các Văn bản hợp nhất (VBHN) tương đương để đánh giá công bằng (vd 39/VBHN-VPQH = hợp nhất Luật Dược).", "Ground-truth — ")
b("chọn Hit@k và MRR — phù hợp bài toán truy xuất, tương ứng Precision/Recall ở mức truy hồi tài liệu.", "Độ đo — ")
note("khớp “Tuần 10: 25/05–31/05 — Xây bộ câu hỏi kiểm thử; kịch bản đánh giá”. Trạng thái: Hoàn thành.")

# ===== TUẦN 11 =====
week("Tuần 11 (2/6 – 8/6): Thực nghiệm đánh giá — phân tích, đánh giá")
b("Tiến hành thực nghiệm, tính các chỉ số và lập bảng/biểu đồ so sánh.")
fig(os.path.join(FIGS, "retrieval_comparison.png"), "Hình 7. So sánh chất lượng truy xuất giữa các cấu hình", 15)
para("Phân tích từng bước:", italic=True)
b("Hit@3 tăng 0,556 → 0,778 (tách từ nhất quán giúp khớp embedding).", "Tách từ pyvi — ")
b("Hit@1 tăng 0,389 → 0,500 (cross-encoder xếp hạng tinh hơn).", "Reranker — ")
b("cắt 500 ký tự sinh gấp ~2,3 lần số đoạn nhưng kém hơn → chọn cắt theo “Điều”.", "Kích thước chunk — ")
b("đạt MRR cao nhất 0,722 và Hit@10 0,944–1,000.", "Hybrid + Reranker — ")
note("khớp “Tuần 11–12: 01/06–14/06 — Thực nghiệm đánh giá, tính Precision/Recall; bảng số liệu, biểu đồ”. Trạng thái: Đang thực hiện.")

# ===== TUẦN 12 =====
week("Tuần 12 (9/6 – 15/6): Hoàn thiện giao diện và báo cáo — phân tích")
b("Hoàn thiện giao diện web React (lịch sử hội thoại, hiển thị nguồn, chạy offline); viết báo cáo theo chuẩn trình bày của trường.")
para("Phân tích/đánh giá: hệ thống đã hoàn chỉnh đầu–cuối (dữ liệu → chỉ mục → truy xuất hybrid + reranker → sinh PhoGPT → web). Sản phẩm và báo cáo sẵn sàng để tổng hợp, hoàn thiện và chuẩn bị bảo vệ.", italic=True)
note("khớp “Tuần 11–12 / 13–14 — hoàn thiện sản phẩm và báo cáo”. Trạng thái: Đang thực hiện.")

try:
    z = doc.settings.element.find(qn('w:zoom'))
    if z is not None and z.get(qn('w:percent')) is None: z.set(qn('w:percent'), '100')
except Exception: pass

doc.save(OUT)
try: print("Da luu:", OUT)
except Exception: pass
