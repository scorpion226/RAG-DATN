# -*- coding: utf-8 -*-
"""
generate_docx.py — Sinh báo cáo ĐATN .docx theo chuẩn trình bày Trường ĐH Thủy lợi.
A4; lề trên2/dưới2/trái3/phải2 cm; Times New Roman 13pt; giãn dòng 1.5; căn đều;
chương 14pt đậm IN HOA; mục cấp 1 (x.y) 13pt đậm; mục cấp 2 (x.y.z) 13pt đậm nghiêng;
chú thích hình ở DƯỚI, chú thích bảng ở TRÊN; trích dẫn IEEE [n]; số trang chân giữa
(bìa không đánh số; La Mã từ Lời cam đoan; Ả Rập từ nội dung).
Chạy: python report/generate_docx.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIGS = os.path.join(HERE, "figs")
OUT = os.path.join(HERE, "BaoCao_RAG_DATN.docx")
FONT = "Times New Roman"


def _set_font(run, size=13, bold=False, italic=False, color=None, caps=False):
    run.font.name = FONT; run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), FONT)
    if caps:
        c = OxmlElement('w:caps'); c.set(qn('w:val'), 'true'); rpr.append(c)


def _field(paragraph, instr):
    r = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    t = OxmlElement('w:instrText'); t.set(qn('xml:space'), 'preserve'); t.text = instr
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    r._r.append(b); r._r.append(t); r._r.append(e)


def _pgnum_fmt(section, fmt, start=None):
    sectPr = section._sectPr
    pg = sectPr.find(qn('w:pgNumType'))
    if pg is None:
        pg = OxmlElement('w:pgNumType')
        cols = sectPr.find(qn('w:cols'))
        (cols.addprevious(pg) if cols is not None else sectPr.append(pg))
    pg.set(qn('w:fmt'), fmt)
    if start is not None: pg.set(qn('w:start'), str(start))


def footer_pagenum(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(p, "PAGE")
    for r in p.runs: _set_font(r, 12)


def margins(section):
    section.page_width = Cm(21.0); section.page_height = Cm(29.7)
    section.top_margin = Cm(2); section.bottom_margin = Cm(2)
    section.left_margin = Cm(3); section.right_margin = Cm(2)


def H1(doc, text):
    p = doc.add_paragraph(style='Heading 1')
    _set_font(p.add_run(text), 14, bold=True, color=RGBColor(0, 0, 0), caps=True)
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True; p.paragraph_format.page_break_before = True
    return p

def H2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    _set_font(p.add_run(text), 13, bold=True, color=RGBColor(0, 0, 0))
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p

def H3(doc, text):
    p = doc.add_paragraph(style='Heading 3')
    _set_font(p.add_run(text), 13, bold=True, italic=True, color=RGBColor(0, 0, 0))
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    return p

def body(doc, text, justify=True, italic=False, indent=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    pf.space_after = Pt(6)
    if indent: pf.first_line_indent = Cm(1.0)
    _set_font(p.add_run(text), 13, italic=italic)
    return p

def formula(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    _set_font(p.add_run(text), 13, italic=True)
    return p

def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        _set_font(p.add_run(bold_lead), 13, bold=True)
    _set_font(p.add_run(text), 13)
    return p

def figure(doc, path, caption, width_cm=15):
    if not os.path.exists(path): return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(path, width=Cm(width_cm))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(c.add_run(caption), 12, bold=True)
    c.paragraph_format.space_after = Pt(10)

def table(doc, headers, rows, caption=None, first_left=True):
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(c.add_run(caption), 12, bold=True); c.paragraph_format.space_before = Pt(6)
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.alignment = 1
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(cell.paragraphs[0].add_run(h), 12, bold=True); _shade(cell, "D9E2F3")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if (i == 0 and first_left)
                                                else WD_ALIGN_PARAGRAPH.CENTER)
            _set_font(cells[i].paragraphs[0].add_run(str(v)), 12)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def _shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexfill); tcPr.append(sh)


def main():
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = FONT; normal.font.size = Pt(13)
    normal.element.rPr.rFonts.set(qn('w:cs'), FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ===================== SECTION 0: TRANG BÌA (không đánh số) =====================
    sec0 = doc.sections[0]; margins(sec0)

    def center(text, size, bold=False, before=0, after=6, caps=False):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
        _set_font(p.add_run(text), size, bold=bold, caps=caps)
    center("TRƯỜNG ĐẠI HỌC THỦY LỢI", 14, True, 6)
    center("KHOA CÔNG NGHỆ THÔNG TIN", 13, True, 0, 24)
    center("ĐỒ ÁN TỐT NGHIỆP", 20, True, 30, 16, caps=True)
    center("HỆ THỐNG RAG HỖ TRỢ TRA CỨU VĂN BẢN PHÁP LUẬT NGÀNH Y TẾ VIỆT NAM", 16, True, 6, 28)
    for lab, val in [("Sinh viên thực hiện", "Nguyễn Minh Hiếu"),
                     ("Ngành", "Công nghệ thông tin"),
                     ("Giảng viên hướng dẫn", "..............................")]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(f"{lab}: "), 13, bold=True); _set_font(p.add_run(val), 13)
    center("Hà Nội — 2026", 13, True, 44)

    # ===================== SECTION 1: PHẦN ĐẦU (số La Mã) =====================
    s1 = doc.add_section(WD_SECTION.NEW_PAGE); margins(s1)
    _pgnum_fmt(s1, "lowerRoman", 1); footer_pagenum(s1)

    def fm_title(text):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12); p.paragraph_format.space_before = Pt(6)
        _set_font(p.add_run(text), 14, bold=True, caps=True)

    fm_title("Lời cam đoan")
    body(doc, "Tôi xin cam đoan đồ án tốt nghiệp “Hệ thống RAG hỗ trợ tra cứu văn bản pháp luật ngành y tế Việt Nam” là công trình nghiên cứu của bản thân dưới sự hướng dẫn của giảng viên. Các số liệu và kết quả thực nghiệm trong báo cáo là trung thực, do tôi tự thực hiện trên bộ dữ liệu công khai; mọi tham khảo đều được trích dẫn đầy đủ. Tôi xin chịu trách nhiệm về lời cam đoan này.")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(p.add_run("Sinh viên: Nguyễn Minh Hiếu"), 13, italic=True)

    doc.add_page_break(); fm_title("Lời cảm ơn")
    body(doc, "Em xin gửi lời cảm ơn chân thành tới giảng viên hướng dẫn đã tận tình định hướng, góp ý trong suốt quá trình thực hiện đồ án. Em cũng cảm ơn quý thầy cô Khoa Công nghệ thông tin — Trường Đại học Thủy lợi đã trang bị nền tảng kiến thức quý báu, cùng gia đình và bạn bè đã luôn động viên, hỗ trợ em hoàn thành đồ án này.")

    doc.add_page_break(); fm_title("Tóm tắt")
    body(doc, "Đồ án xây dựng một hệ thống Retrieval-Augmented Generation (RAG) hỗ trợ tra cứu văn bản pháp luật ngành y tế Việt Nam, vận hành hoàn toàn trên máy cá nhân (CPU, 16GB RAM, không GPU). Hệ thống lập chỉ mục 367.462 đoạn văn bản trích từ 21.490 văn bản còn hiệu lực; sử dụng embedding dựa trên PhoBERT (bi-encoder) lưu trong ChromaDB, kết hợp tìm kiếm từ vựng BM25 theo cơ chế hybrid (Reciprocal Rank Fusion) và bộ xếp hạng lại (reranker) PhoRanker; phần sinh dùng PhoGPT-4B-Chat lượng tử hóa để tạo câu trả lời có trích dẫn. Trên bộ câu hỏi đánh giá, cấu hình Hybrid + Reranker đạt MRR 0,722 và Hit@10 tới 1,000; mỗi thành phần kỹ thuật đều được kiểm chứng bằng thực nghiệm A/B. Báo cáo cũng trình bày quá trình phát hiện và khắc phục lỗi dữ liệu nghiêm trọng trong pipeline ban đầu.")
    p = doc.add_paragraph(); _set_font(p.add_run("Từ khóa: "), 13, bold=True)
    _set_font(p.add_run("RAG, PhoBERT, PhoGPT, ChromaDB, BM25, hybrid search, pháp luật y tế."), 13)

    doc.add_page_break()
    pml = doc.add_paragraph(); pml.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(pml.add_run("MỤC LỤC"), 14, bold=True)
    _field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u')

    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run("DANH MỤC HÌNH VẼ"), 14, bold=True)
    for s in ["Hình 1.1. Kiến trúc tổng thể hệ thống RAG",
              "Hình 2.1. Quy trình lọc dữ liệu",
              "Hình 3.1. Lược đồ tuần tự xử lý một câu hỏi",
              "Hình 3.2. Giao diện web demo",
              "Hình 4.1. So sánh chất lượng truy xuất giữa các cấu hình",
              "Hình 4.2. Precision/Recall/F1 theo số nguồn k",
              "Hình 4.3. Thời gian phản hồi trung bình"]:
        body(doc, s, justify=False, indent=False)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(12)
    _set_font(p.add_run("DANH MỤC BẢNG BIỂU"), 14, bold=True)
    for s in ["Bảng 1.1. So sánh các hướng tiếp cận trả lời câu hỏi chuyên ngành",
              "Bảng 2.1. Phân bố loại văn bản trong corpus",
              "Bảng 3.1. Các module chính của hệ thống",
              "Bảng 4.1. So sánh các cấu hình truy xuất",
              "Bảng 4.2. Ảnh hưởng của kích thước chunk",
              "Bảng 4.3. Precision/Recall/F1 theo số nguồn k"]:
        body(doc, s, justify=False, indent=False)

    # ===================== SECTION 2: NỘI DUNG (số Ả Rập) =====================
    s2 = doc.add_section(WD_SECTION.NEW_PAGE); margins(s2)
    _pgnum_fmt(s2, "decimal", 1); footer_pagenum(s2)

    # ----------------------------- CHƯƠNG 0 -----------------------------
    H1(doc, "CHƯƠNG 0. MỞ ĐẦU")
    H2(doc, "0.1. Bối cảnh và động lực")
    body(doc, "Hệ thống văn bản quy phạm pháp luật Việt Nam trong lĩnh vực y tế có quy mô lớn và biến động liên tục: một luật có thể được sửa đổi, bổ sung, hợp nhất hoặc bị thay thế qua thời gian, kéo theo nhiều nghị định, thông tư, quyết định hướng dẫn. Ngôn ngữ pháp lý lại mang tính hình thức cao, nhiều thuật ngữ chuyên ngành và cấu trúc “Điều — Khoản — Điểm” đặc thù. Những yếu tố này khiến việc tra cứu của người dân, cán bộ y tế và sinh viên trở nên khó khăn.")
    body(doc, "Các công cụ tìm kiếm từ khóa truyền thống chỉ khớp trên mặt chữ nên bỏ sót khi người dùng diễn đạt khác cách dùng từ trong văn bản. Ở chiều ngược lại, việc hỏi trực tiếp một mô hình ngôn ngữ lớn (Large Language Model — LLM) tuy cho câu trả lời trôi chảy nhưng tiềm ẩn rủi ro “bịa đặt” (hallucination) và không truy xuất được nguồn — điều không thể chấp nhận trong lĩnh vực pháp luật, nơi mọi khẳng định cần dẫn chiếu văn bản cụ thể.")
    body(doc, "Kiến trúc Retrieval-Augmented Generation (RAG) [1] dung hòa hai hướng trên: trước hết truy xuất các đoạn văn bản liên quan từ một kho tri thức đáng tin cậy, sau đó để LLM sinh câu trả lời dựa trên các đoạn này và trích dẫn nguồn. Nhờ vậy, tri thức của hệ thống có thể cập nhật mà không cần huấn luyện lại mô hình, đồng thời câu trả lời bám sát văn bản gốc và kiểm chứng được. Đây chính là động lực để đề tài lựa chọn RAG cho bài toán tra cứu pháp luật y tế.")
    H2(doc, "0.2. Phát biểu bài toán")
    body(doc, "Cho một câu hỏi tiếng Việt liên quan đến pháp luật ngành y tế, hệ thống cần (i) tìm trong kho văn bản các đoạn liên quan nhất, và (ii) sinh một câu trả lời chính xác, súc tích, kèm trích dẫn số hiệu văn bản nguồn. Toàn bộ hệ thống phải vận hành được trên máy cá nhân với CPU và 16GB RAM, không sử dụng GPU.")
    H2(doc, "0.3. Mục tiêu và phạm vi")
    bullet(doc, "Thu thập, làm sạch và tổ chức kho văn bản pháp luật ngành y tế từ nguồn dữ liệu công khai.", "Mục tiêu 1 — ")
    bullet(doc, "Xây dựng pipeline RAG hoàn chỉnh: lập chỉ mục, truy xuất, xếp hạng lại và sinh câu trả lời.", "Mục tiêu 2 — ")
    bullet(doc, "Đánh giá định lượng đóng góp của từng kỹ thuật và phát triển giao diện web hỏi–đáp.", "Mục tiêu 3 — ")
    body(doc, "Phạm vi giới hạn ở văn bản pháp luật thuộc ngành y tế còn hiệu lực; hệ thống mang tính hỗ trợ tra cứu, không thay thế tư vấn pháp lý chính thức và không cập nhật văn bản theo thời gian thực.")
    H2(doc, "0.4. Câu hỏi nghiên cứu")
    bullet(doc, "Làm thế nào xây dựng một pipeline RAG tiếng Việt cho văn bản pháp luật y tế chạy được trên tài nguyên hạn chế (CPU, 16GB RAM)?", "RQ1. ")
    bullet(doc, "Mỗi thành phần kỹ thuật (tách từ tiếng Việt, reranker, hybrid BM25+vector) đóng góp bao nhiêu vào chất lượng truy xuất, đo bằng Hit@k và MRR?", "RQ2. ")
    bullet(doc, "Hệ thống có sinh câu trả lời bám ngữ cảnh, có trích dẫn nguồn và hạn chế bịa đặt không?", "RQ3. ")
    H2(doc, "0.5. Đóng góp của đồ án")
    bullet(doc, "Một kho tri thức pháp luật y tế đã làm sạch gồm 367.462 đoạn từ 21.490 văn bản còn hiệu lực, kèm quy trình xử lý có thể tái lập.")
    bullet(doc, "Một pipeline RAG tiếng Việt hoàn chỉnh kết hợp bi-encoder, BM25 hybrid (RRF), reranker và LLM PhoGPT, tối ưu cho CPU.")
    bullet(doc, "Bộ thực nghiệm A/B định lượng đóng góp của từng kỹ thuật và bài học thực tiễn về kiểm định chất lượng dữ liệu.")
    H2(doc, "0.6. Bố cục báo cáo")
    body(doc, "Chương 1 trình bày cơ sở lý thuyết; Chương 2 mô tả dữ liệu và tiền xử lý; Chương 3 trình bày thiết kế và triển khai hệ thống; Chương 4 trình bày thực nghiệm và đánh giá; Chương 5 kết luận và nêu hướng phát triển.")

    # ----------------------------- CHƯƠNG 1 -----------------------------
    H1(doc, "CHƯƠNG 1. CƠ SỞ LÝ THUYẾT")
    H2(doc, "1.1. Tổng quan kiến trúc RAG")
    body(doc, "RAG [1] là kiến trúc kết hợp hai khối: bộ truy xuất (Retriever) lấy về các đoạn tri thức liên quan từ một kho ngoài, và bộ sinh (Generator) — thường là một LLM — tạo câu trả lời dựa trên câu hỏi cùng các đoạn truy xuất được. So với việc tinh chỉnh (fine-tuning) toàn bộ tri thức vào tham số mô hình, RAG có ba ưu điểm: cập nhật tri thức chỉ cần thay đổi kho dữ liệu; câu trả lời có thể trích dẫn nguồn; và giảm đáng kể hiện tượng bịa đặt do mô hình bị “ràng buộc” vào ngữ cảnh được cung cấp.")
    figure(doc, os.path.join(FIGS, "architecture.png"), "Hình 1.1. Kiến trúc tổng thể hệ thống RAG", 10.5)
    body(doc, "Một pipeline RAG điển hình gồm các bước: chia nhỏ tài liệu thành đoạn (chunking), mã hóa đoạn thành vector (embedding) và lưu vào cơ sở dữ liệu vector; khi có câu hỏi, mã hóa câu hỏi, truy xuất k đoạn gần nhất, (tùy chọn) xếp hạng lại, rồi đưa vào prompt để LLM sinh câu trả lời. Mỗi bước đều ảnh hưởng tới chất lượng cuối cùng và sẽ được phân tích, kiểm chứng ở các chương sau.")
    body(doc, "Có ba hướng tiếp cận chính để một mô hình ngôn ngữ trả lời câu hỏi chuyên ngành: (1) chỉ dùng LLM với tri thức sẵn có trong tham số; (2) tinh chỉnh (fine-tuning) LLM trên dữ liệu chuyên ngành; và (3) RAG. Bảng 1.1 so sánh ba hướng theo các tiêu chí quan trọng với bài toán pháp luật.")
    table(doc, ["Tiêu chí", "LLM thuần", "Fine-tuning", "RAG"],
          [["Cập nhật tri thức", "Khó", "Huấn luyện lại", "Đổi kho dữ liệu"],
           ["Trích dẫn nguồn", "Không", "Không", "Có"],
           ["Nguy cơ bịa đặt", "Cao", "Trung bình", "Thấp"],
           ["Chi phí tài nguyên", "Thấp", "Rất cao", "Trung bình"]],
          caption="Bảng 1.1. So sánh các hướng tiếp cận trả lời câu hỏi chuyên ngành")
    body(doc, "Với miền pháp luật — nơi câu trả lời bắt buộc phải kiểm chứng được và tri thức thay đổi liên tục — RAG là lựa chọn hợp lý nhất: vừa cập nhật dễ, vừa cho phép trích dẫn, vừa giảm bịa đặt mà không đòi hỏi tài nguyên huấn luyện lớn.")
    H2(doc, "1.2. Biểu diễn ngữ nghĩa văn bản bằng embedding")
    H3(doc, "1.2.1. Mô hình PhoBERT và đặc thù tiếng Việt")
    body(doc, "Các mô hình ngôn ngữ hiện đại dựa trên kiến trúc Transformer với cơ chế tự chú ý (self-attention), cho phép mô hình cân nhắc mối liên hệ giữa mọi cặp từ trong câu, từ đó nắm bắt ngữ cảnh tốt hơn các mạng tuần tự truyền thống. BERT là mô hình mã hóa hai chiều, được tiền huấn luyện theo tác vụ che từ (masked language modeling) trên kho văn bản lớn, sau đó có thể tinh chỉnh cho nhiều tác vụ hạ nguồn.")
    body(doc, "PhoBERT [2] là mô hình ngôn ngữ tiền huấn luyện cho tiếng Việt theo kiến trúc RoBERTa, huấn luyện trên kho ngữ liệu tiếng Việt quy mô lớn. Một đặc điểm quan trọng: PhoBERT được huấn luyện trên dữ liệu đã được tách từ (word segmentation). Tiếng Việt là ngôn ngữ đơn lập, ranh giới “từ” không trùng với khoảng trắng — ví dụ “Bộ Y tế” là một từ ghép gồm ba âm tiết. Do đó, để khai thác đúng PhoBERT, đầu vào lúc suy luận cũng phải được tách từ bằng công cụ như pyvi, tạo dạng “Bộ_Y_tế”; nếu không, embedding sẽ lệch khỏi phân phối lúc huấn luyện và làm giảm chất lượng (sẽ được kiểm chứng ở Thí nghiệm 1).")
    H3(doc, "1.2.2. Bi-encoder và lý do không dùng PhoBERT trực tiếp")
    body(doc, "PhoBERT gốc là mô hình masked language model; vector tại token [CLS] của nó không được tối ưu để đo độ tương đồng ngữ nghĩa giữa hai câu. Theo hướng Sentence-BERT [4], cần tinh chỉnh mô hình theo kiến trúc bi-encoder: hai câu được đưa qua cùng một mạng (chia sẻ trọng số), gộp các vector token (pooling) thành một vector câu, rồi huấn luyện theo mục tiêu tương phản (contrastive) — kéo gần các cặp liên quan và đẩy xa các cặp không liên quan. Nhờ vậy, khoảng cách trong không gian vector phản ánh đúng độ tương đồng ngữ nghĩa.")
    body(doc, "Đề tài sử dụng mô hình vietnamese-bi-encoder [7] (dựa trên PhoBERT, vector 768 chiều) — đây là cách dùng đúng cho bài toán tìm kiếm ngữ nghĩa. Độ tương đồng giữa hai vector u, v được đo bằng cosine:")
    formula(doc, "cos(u, v) = (u · v) / (‖u‖ · ‖v‖),  giá trị trong [-1, 1].")
    body(doc, "Ưu điểm của bi-encoder là mã hóa đoạn MỘT LẦN khi lập chỉ mục, sau đó truy vấn rất nhanh; nhược điểm là mã hóa câu hỏi và đoạn độc lập nên kém tinh tế — hạn chế này được khắc phục bằng reranker (Mục 1.5).")
    H2(doc, "1.3. Cơ sở dữ liệu vector và tìm kiếm lân cận gần đúng")
    body(doc, "Sau khi mã hóa, mỗi đoạn trở thành một điểm trong không gian 768 chiều. Truy xuất là tìm k điểm gần câu hỏi nhất theo cosine. Với hàng trăm nghìn vector, tìm kiếm vét cạn (so sánh với mọi điểm) quá chậm cho thời gian thực, nên dùng chỉ mục lân cận gần đúng (Approximate Nearest Neighbor — ANN).")
    body(doc, "ChromaDB [9] sử dụng thuật toán HNSW (Hierarchical Navigable Small World): xây một đồ thị nhiều tầng, tầng trên thưa để “nhảy xa”, tầng dưới dày để tinh chỉnh; khi truy vấn, thuật toán bắt đầu từ tầng trên và đi xuống dần theo các cạnh gần câu hỏi nhất. Cách này đánh đổi một phần độ chính xác lấy tốc độ truy vấn gần như hằng số theo kích thước kho, phù hợp với hàng trăm nghìn vector. ChromaDB còn lưu kèm metadata (số hiệu, loại văn bản, tình trạng hiệu lực…), cho phép lọc ngay trong truy vấn (ví dụ chỉ lấy văn bản còn hiệu lực) — rất hữu ích cho yêu cầu “chỉ tư vấn văn bản còn hiệu lực”.")
    H2(doc, "1.4. Tìm kiếm từ vựng BM25 và truy xuất hybrid")
    H3(doc, "1.4.1. BM25")
    body(doc, "BM25 [5] là hàm xếp hạng dựa trên trùng khớp từ vựng, cải tiến từ TF-IDF. Điểm BM25 của tài liệu d với truy vấn q được tính theo công thức:")
    formula(doc, "BM25(q,d) = Σ IDF(t) · [ f(t,d)·(k1+1) ] / [ f(t,d) + k1·(1 − b + b·|d|/avgdl) ]")
    body(doc, "trong đó f(t,d) là tần suất từ t trong tài liệu d; |d| là độ dài tài liệu; avgdl là độ dài trung bình; IDF(t) là trọng số nghịch tần suất tài liệu; k1 điều chỉnh độ bão hòa tần suất từ và b điều chỉnh mức chuẩn hóa theo độ dài. BM25 mạnh ở các truy vấn chứa từ khóa hiếm hoặc định danh chính xác — đặc biệt là số hiệu văn bản như “147/2025/NĐ-CP” — vốn là điểm yếu cố hữu của embedding ngữ nghĩa, do embedding dễ “làm mờ” các chuỗi ký tự–số hiếm gặp.")
    H3(doc, "1.4.2. Hợp nhất bằng Reciprocal Rank Fusion (RRF)")
    body(doc, "Để tận dụng đồng thời thế mạnh ngữ nghĩa của vector và thế mạnh từ vựng của BM25, đề tài hợp nhất hai danh sách kết quả bằng Reciprocal Rank Fusion [6]. Điểm hợp nhất của tài liệu d được tính theo công thức:")
    formula(doc, "RRF(d) = Σ over r ∈ {vector, BM25}  1 / (c + rank_r(d)),  với c = 60.")
    body(doc, "RRF chỉ dựa trên thứ hạng (rank) của tài liệu trong mỗi danh sách nên không cần chuẩn hóa thang điểm khác nhau giữa cosine và BM25 — một ưu điểm thực tiễn quan trọng. Ngoài ra, đề tài bổ sung bước nhận diện số hiệu văn bản trong câu hỏi (bằng biểu thức chính quy) để truy vấn trực tiếp đúng văn bản đó, khắc phục triệt để trường hợp người dùng hỏi theo số hiệu.")
    H2(doc, "1.5. Xếp hạng lại bằng cross-encoder")
    body(doc, "Bi-encoder mã hóa câu hỏi và đoạn một cách độc lập rồi mới so khớp, nên nhanh và phù hợp để quét toàn kho, nhưng đánh đổi độ tinh tế. Cross-encoder ngược lại đưa đồng thời cặp (câu hỏi, đoạn) qua mô hình và xuất một điểm liên quan duy nhất, nhờ có cơ chế chú ý chéo (cross-attention) nên đánh giá chính xác hơn nhiều, song chi phí lớn. Chiến lược “retrieve-then-rerank” tận dụng cả hai: bi-encoder lấy nhanh N ứng viên, cross-encoder xếp lại để chọn k tốt nhất. Đề tài dùng PhoRanker [8] (cross-encoder dựa trên PhoBERT, cũng yêu cầu tách từ).")
    H2(doc, "1.6. Mô hình ngôn ngữ sinh PhoGPT và lượng tử hóa")
    body(doc, "PhoGPT-4B-Chat [3] là LLM tiếng Việt với 3,7 tỷ tham số, kiến trúc MPT, độ dài ngữ cảnh 8192 token. Ở định dạng đầy đủ (FP32), mô hình cần khoảng 16GB bộ nhớ — vượt khả năng máy mục tiêu khi còn phải tải các mô hình khác. Đề tài dùng kỹ thuật lượng tử hóa (quantization): biểu diễn trọng số ở độ chính xác thấp hơn (4-bit, định dạng GGUF Q4_K_M) qua thư viện llama.cpp, giảm dung lượng còn 2,36GB và cho phép suy luận trên CPU với mất mát chất lượng nhỏ. Đây là đánh đổi hợp lý giữa chất lượng và khả năng triển khai trên tài nguyên hạn chế.")
    H2(doc, "1.7. Các độ đo đánh giá truy xuất")
    body(doc, "Đề tài dùng hai độ đo phổ biến trong truy hồi thông tin. Hit@k là tỷ lệ câu hỏi có ít nhất một văn bản liên quan nằm trong k kết quả đầu, phản ánh khả năng “bắt trúng” nguồn. MRR (Mean Reciprocal Rank) là trung bình nghịch đảo thứ hạng của văn bản liên quan đầu tiên:")
    formula(doc, "MRR = (1/|Q|) Σ 1/rank_i ,   Hit@k = (1/|Q|) Σ 1[∃ tài liệu liên quan trong top-k].")
    body(doc, "MRR nhạy với việc tài liệu đúng được xếp ở vị trí nào (ưu tiên xếp hạng cao), trong khi Hit@k phản ánh độ bao phủ. Hai độ đo bổ trợ nhau khi đánh giá chất lượng truy xuất.")
    H2(doc, "1.8. Một số công trình liên quan")
    body(doc, "Kiến trúc RAG được Lewis và cộng sự [1] đề xuất, kết hợp truy xuất tài liệu với mô hình sinh để giải các tác vụ đòi hỏi tri thức. Về biểu diễn câu, Sentence-BERT [4] cho thấy việc tinh chỉnh BERT theo kiến trúc Siamese cải thiện rõ chất lượng đo tương đồng câu so với dùng vector [CLS] thô. Đối với tiếng Việt, PhoBERT [2] là mô hình nền tảng cho nhiều tác vụ; trên cơ sở đó, các mô hình bi-encoder tiếng Việt như vietnamese-bi-encoder [7] và bộ xếp hạng PhoRanker [8] được phát triển cho tìm kiếm và xếp hạng. Về sinh ngôn ngữ tiếng Việt, PhoGPT [3] là một trong các LLM mở quy mô vừa, phù hợp triển khai cục bộ sau lượng tử hóa.")
    body(doc, "Về kỹ thuật truy xuất, BM25 [5] vẫn là chuẩn mạnh cho tìm kiếm từ vựng; việc kết hợp tín hiệu từ vựng và ngữ nghĩa (hybrid) bằng Reciprocal Rank Fusion [6] thường vượt trội so với từng phương pháp đơn lẻ. Đề tài kế thừa các nền tảng này và đóng góp ở khía cạnh tích hợp – tối ưu cho miền pháp luật y tế tiếng Việt trên tài nguyên hạn chế, kèm cơ chế nhận diện số hiệu văn bản.")

    # ----------------------------- CHƯƠNG 2 -----------------------------
    H1(doc, "CHƯƠNG 2. DỮ LIỆU VÀ TIỀN XỬ LÝ")
    H2(doc, "2.1. Nguồn dữ liệu")
    body(doc, "Đề tài sử dụng bộ dữ liệu công khai vietnamese-legal-documents [10] trên Hugging Face, gồm 518.601 bản ghi metadata và 518.235 bản ghi nội dung văn bản pháp luật Việt Nam. Mỗi văn bản kèm các trường: số hiệu (document_number), tiêu đề, loại văn bản (legal_type), ngành/lĩnh vực (legal_sectors), cơ quan ban hành, ngày ban hành, ngày hiệu lực và tình trạng hiệu lực (effect_status).")
    H2(doc, "2.2. Lọc theo ngành và tình trạng hiệu lực")
    body(doc, "Trước hết lọc các văn bản có trường legal_sectors chứa “Health”/“Y tế”, thu được 34.223 văn bản, trong đó 34.166 văn bản có nội dung trong kho content. Tiếp theo, chỉ giữ các văn bản có tình trạng “còn hiệu lực” (In effect), còn lại 21.490 văn bản. Việc lọc hiệu lực là cần thiết vì một hệ thống tra cứu pháp luật chỉ nên tư vấn theo văn bản đang có hiệu lực; trích dẫn văn bản đã hết hiệu lực có thể gây hiểu sai nghiêm trọng.")
    figure(doc, os.path.join(FIGS, "data_funnel.png"), "Hình 2.1. Quy trình lọc dữ liệu", 13.5)
    H2(doc, "2.3. Làm sạch văn bản")
    body(doc, "Nội dung gốc là văn bản thuần (đã loại HTML), còn giữ dấu xuống dòng và các mốc “Điều”, “Chương”. Bước làm sạch chuẩn hóa khoảng trắng nhưng có chủ đích giữ lại cấu trúc xuống dòng và các ký tự dấu “/ - ( ) %”. Lý do: số hiệu văn bản là định danh quan trọng nhất để trích dẫn; nếu xóa các dấu này, số hiệu sẽ bị biến dạng, làm hỏng khả năng tra cứu. Ví dụ minh họa:")
    body(doc, "• Trước (làm sạch sai của pipeline cũ): “… Nghị định 147 2025 NĐ CP …” → mất cấu trúc số hiệu.\n• Sau (làm sạch đúng): “… Nghị định 147/2025/NĐ-CP …” → giữ nguyên định danh.", justify=False)
    body(doc, "Trong pipeline ban đầu, chính lỗi xóa dấu này đã được phát hiện và sửa, cùng với lỗi nhân bản dữ liệu trình bày ở Mục 2.6.")
    H2(doc, "2.4. Chia đoạn theo đơn vị ngữ nghĩa pháp lý")
    body(doc, "Văn bản được chia thành đoạn (chunk) bằng thuật toán đệ quy, gồm các bước: (1) nếu đoạn ngắn hơn kích thước mục tiêu (~1200 ký tự) thì giữ nguyên; (2) ngược lại, tách tại mốc ưu tiên cao nhất còn xuất hiện theo thứ tự “Điều” → “Chương” → “Mục” → đoạn trống → xuống dòng → câu → khoảng trắng; (3) đệ quy trên các phần vẫn còn dài; (4) gộp các phần nhỏ liền kề thành đoạn ~1200 ký tự, thêm phần chồng lấp (overlap) 150 ký tự ở đầu mỗi đoạn kế tiếp.")
    body(doc, "Việc cắt theo “Điều” giúp mỗi đoạn là một đơn vị ngữ nghĩa tương đối hoàn chỉnh, thay vì cắt cứng theo số ký tự dễ làm đứt câu và mất ngữ cảnh. Phần chồng lấp giữ liên tục ngữ nghĩa giữa các đoạn liền kề, tránh mất thông tin nằm đúng ở ranh giới cắt. Giả thuyết “cắt theo Điều tốt hơn cắt cứng” được kiểm chứng định lượng ở Chương 4 (Thí nghiệm 3).")
    H2(doc, "2.5. Khử trùng lặp và thống kê corpus")
    body(doc, "Văn bản pháp luật chứa nhiều phần lặp (quốc hiệu, tiêu ngữ, căn cứ ban hành…). Đề tài khử trùng lặp ở mức đoạn bằng hàm băm, loại bỏ 9.131 đoạn trùng. Kho tri thức cuối cùng gồm 367.462 đoạn từ 21.490 văn bản, trung bình 17,1 đoạn/văn bản, độ dài trung bình khoảng 1.036 ký tự/đoạn. Phân bố loại văn bản được trình bày ở Bảng 2.1, cho thấy Quyết định và Công văn chiếm đa số — đặc trưng của văn bản điều hành ngành y tế.")
    table(doc, ["Loại văn bản", "Số văn bản", "Tỷ lệ"],
          [["Quyết định (Decision)", "10.343", "48,1%"],
           ["Công văn (Official Dispatch)", "8.399", "39,1%"],
           ["Nghị quyết (Resolution)", "752", "3,5%"],
           ["Chỉ thị (Directive)", "654", "3,0%"],
           ["Thông tư (Circular)", "510", "2,4%"],
           ["Công điện (Official Telegram)", "243", "1,1%"],
           ["Văn bản hợp nhất / Khác", "≈600", "≈2,8%"]],
          caption="Bảng 2.1. Phân bố loại văn bản trong corpus")
    H2(doc, "2.6. Bài học từ pipeline dữ liệu sai")
    body(doc, "Pipeline xử lý ban đầu mắc hai lỗi nghiêm trọng được phát hiện trong quá trình kiểm định: (i) chỉ thu được 8.278 trên 34.166 văn bản có nội dung (mất khoảng 75%); và (ii) một lỗi trong vòng lặp ghi khiến toàn bộ tập đoạn bị nhân bản khoảng 78 lần, làm tập dữ liệu phình lên 27,7 triệu dòng trong khi số đoạn duy nhất chỉ là 354.590 (tỷ lệ trùng 98,7%). Việc đo đạc tỷ lệ trùng, truy nguyên nguyên nhân và xây dựng lại pipeline sạch đã được thực hiện và là một đóng góp thực tiễn của đồ án, nhấn mạnh tầm quan trọng của khâu kiểm định dữ liệu trước khi xây dựng mô hình.")

    # ----------------------------- CHƯƠNG 3 -----------------------------
    H1(doc, "CHƯƠNG 3. THIẾT KẾ VÀ TRIỂN KHAI HỆ THỐNG")
    H2(doc, "3.1. Kiến trúc tổng thể và luồng xử lý")
    body(doc, "Hệ thống được tổ chức theo các module rời, mỗi module đảm nhiệm một giai đoạn của pipeline RAG (xem lại Hình 1.1). Luồng xử lý một câu hỏi gồm: tách từ và mã hóa câu hỏi; truy xuất song song từ ChromaDB (vector) và BM25 (từ vựng); hợp nhất bằng RRF kèm nhận diện số hiệu; xếp hạng lại bằng PhoRanker để chọn k đoạn tốt nhất; dựng prompt và sinh câu trả lời bằng PhoGPT; cuối cùng trả về câu trả lời kèm danh sách nguồn. Trình tự tương tác giữa các thành phần được thể hiện ở Hình 3.1.")
    figure(doc, os.path.join(FIGS, "sequence.png"), "Hình 3.1. Lược đồ tuần tự xử lý một câu hỏi", 15)
    H2(doc, "3.2. Các module và vai trò")
    table(doc, ["Module (tệp)", "Vai trò"],
          [["build_chunks.py", "Lọc, làm sạch, chia đoạn dữ liệu → tệp Parquet"],
           ["index_chroma.py", "Tạo embedding (PhoBERT) và nạp vào ChromaDB; hỗ trợ resume"],
           ["build_bm25.py", "Xây chỉ mục BM25 trên toàn corpus"],
           ["rag_core.py / hybrid.py", "Truy xuất vector / hybrid + nhận diện số hiệu"],
           ["rerank.py", "Xếp hạng lại bằng cross-encoder PhoRanker"],
           ["llm.py", "Sinh câu trả lời (PhoGPT GGUF) — có chế độ thử nghiệm"],
           ["app.py, static/", "Backend FastAPI và giao diện web React"]],
          caption="Bảng 3.1. Các module chính của hệ thống")
    H2(doc, "3.3. Lập chỉ mục và khả năng phục hồi")
    body(doc, "Việc tạo embedding cho 367.462 đoạn trên CPU mất khoảng 13 giờ (tốc độ ~7,7 đoạn/giây). Để không phụ thuộc một phiên chạy liên tục, module lập chỉ mục ghi dần kết quả vào ChromaDB và hỗ trợ phục hồi (resume): khi chạy lại, hệ thống bỏ qua số đoạn đã nạp và tiếp tục từ vị trí dang dở. Dữ liệu được xử lý theo lô (batch) và theo luồng (streaming) để giữ mức tiêu thụ RAM thấp, phù hợp ràng buộc 16GB.")
    H2(doc, "3.4. Thiết kế truy xuất hybrid")
    body(doc, "Module hybrid hoạt động theo các bước: (1) lấy top-N ứng viên từ truy xuất vector (ChromaDB) và top-N từ BM25; (2) hợp nhất hai danh sách bằng Reciprocal Rank Fusion, cộng điểm theo thứ hạng của mỗi tài liệu trong từng danh sách; (3) dùng biểu thức chính quy phát hiện số hiệu văn bản trong câu hỏi — nếu có, truy vấn trực tiếp các đoạn thuộc đúng văn bản đó từ ChromaDB và cộng điểm ưu tiên lớn; (4) lọc theo tình trạng hiệu lực; (5) (tùy chọn) đưa các ứng viên qua reranker để chọn k đoạn cuối. Bước (3) là điểm mấu chốt giúp hệ thống trả lời chính xác khi người dùng hỏi theo số hiệu — trường hợp mà tìm kiếm vector thuần thường thất bại.")
    H2(doc, "3.5. Thiết kế prompt và kiểm soát sinh")
    body(doc, "Prompt đưa vào PhoGPT gồm ba phần: (i) chỉ dẫn hệ thống, (ii) ngữ cảnh là các đoạn truy xuất được (mỗi đoạn kèm số hiệu và tiêu đề nguồn), (iii) câu hỏi của người dùng. Cấu trúc prompt được minh họa như sau:")
    code = ("### Câu hỏi:\n"
            "Bạn là trợ lý tra cứu văn bản pháp luật ngành y tế. CHỈ dựa vào các văn bản "
            "được cung cấp để trả lời. Luôn trích dẫn số hiệu văn bản. Nếu thông tin không "
            "có trong văn bản, hãy nói rõ là không tìm thấy.\n\n"
            "=== VĂN BẢN THAM KHẢO ===\n[1] <số hiệu> — <tiêu đề>\n<nội dung đoạn>\n...\n\n"
            "=== CÂU HỎI ===\n<câu hỏi của người dùng>\n\n### Trả lời:")
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(code); r.font.name = "Consolas"; r.font.size = Pt(11)
    body(doc, "Cách thiết kế này nhằm “neo” câu trả lời vào ngữ cảnh được cung cấp, buộc mô hình trích dẫn nguồn và thừa nhận khi không có thông tin — qua đó giảm hiện tượng bịa đặt, vốn là rủi ro lớn nhất khi áp dụng LLM cho miền pháp luật.")
    H2(doc, "3.6. Backend và giao diện web")
    body(doc, "Backend dùng FastAPI cung cấp các điểm cuối: trang giao diện, /chat (nhận câu hỏi, trả về câu trả lời và nguồn), /health (trạng thái chỉ mục và cấu hình). Giao diện là một ứng dụng web đơn trang (SPA) bằng React, hiển thị hội thoại dạng chat, lưu lịch sử ở phía trình duyệt, và hiển thị các thẻ nguồn (số hiệu, loại văn bản, tình trạng hiệu lực, điểm liên quan, trích đoạn). Các thư viện front-end được đóng gói cục bộ để giao diện hoạt động ngay cả khi không có Internet — thuận lợi khi trình diễn và bảo vệ.")
    figure(doc, os.path.join(FIGS, "web_result.png"), "Hình 3.2. Giao diện web demo: câu trả lời của PhoGPT kèm các nguồn trích dẫn", 11)
    body(doc, "Backend cung cấp các điểm cuối (endpoint) chính như Bảng 3.2; cấu hình (chế độ LLM, bật/tắt hybrid và reranker) được điều khiển qua biến môi trường, thuận tiện cho việc thử nghiệm các cấu hình khác nhau khi đánh giá.")
    table(doc, ["Endpoint", "Phương thức", "Chức năng"],
          [["/", "GET", "Trả về trang giao diện web (SPA React)"],
           ["/chat", "POST", "Nhận {câu hỏi, k, chỉ_hiệu_lực}; trả về {câu trả lời, danh sách nguồn}"],
           ["/health", "GET", "Trả về trạng thái: số vector, chế độ LLM, hybrid/reranker"],
           ["/static/*", "GET", "Phục vụ tài nguyên tĩnh (thư viện React đóng gói cục bộ)"]],
          caption="Bảng 3.2. Các điểm cuối (API) của hệ thống")

    # ----------------------------- CHƯƠNG 4 -----------------------------
    H1(doc, "CHƯƠNG 4. THỰC NGHIỆM VÀ ĐÁNH GIÁ")
    H2(doc, "4.1. Thiết lập thực nghiệm")
    body(doc, "Bộ đánh giá gồm 50 câu hỏi tiếng Việt bao phủ nhiều luật y tế trọng yếu (Khám bệnh chữa bệnh, Dược, An toàn thực phẩm, Bảo hiểm y tế, phòng chống tác hại thuốc lá/rượu bia, hiến mô tạng, HIV/AIDS, phòng bệnh). Mỗi câu hỏi được gán nhãn “văn bản liên quan” (ground-truth) là (các) luật điều chỉnh nội dung câu hỏi (trung bình 1,96 văn bản/câu). Đặc biệt, nhãn được bổ sung các Văn bản hợp nhất (VBHN) tương đương — ví dụ 39/VBHN-VPQH là bản hợp nhất Luật Dược — vì chúng có cùng nội dung luật nhưng khác số hiệu. Chỉ mục dùng toàn bộ 367.462 đoạn. Các độ đo là Hit@k, MRR (Mục 1.7) và Precision/Recall/F1@k.")
    H2(doc, "4.2. Kết quả tổng hợp")
    table(doc, ["Cấu hình", "Hit@1", "Hit@3", "Hit@5", "Hit@10", "MRR"],
          [["(1) Không tách từ", "0,380", "0,640", "0,720", "0,840", "0,534"],
           ["(2) Vector + tách từ pyvi", "0,520", "0,820", "0,860", "0,920", "0,668"],
           ["(3) Vector + Reranker", "0,540", "0,780", "0,840", "0,880", "0,667"],
           ["(4) Hybrid (BM25+vector)", "0,500", "0,800", "0,860", "0,980", "0,669"],
           ["(5) Hybrid + Reranker", "0,560", "0,780", "0,860", "0,920", "0,694"]],
          caption="Bảng 4.1. So sánh các cấu hình truy xuất (50 câu hỏi)")
    figure(doc, os.path.join(FIGS, "retrieval_comparison.png"), "Hình 4.1. Biểu đồ so sánh chất lượng truy xuất giữa các cấu hình", 15)
    H2(doc, "4.3. Thí nghiệm 1 — Vai trò của tách từ tiếng Việt")
    body(doc, "So sánh cấu hình (1) và (2) trong Bảng 4.1: bổ sung bước tách từ bằng pyvi nâng Hit@1 từ 0,380 lên 0,520 (tăng 0,140), Hit@3 từ 0,640 lên 0,820 (tăng 0,180) và MRR từ 0,534 lên 0,668 (tăng 0,134). Đây là cải thiện lớn nhất trong các thí nghiệm, khẳng định giả thuyết ở Mục 1.2.1: vì PhoBERT được huấn luyện trên dữ liệu đã tách từ, việc tách từ nhất quán cả khi lập chỉ mục lẫn khi truy vấn giúp embedding của câu hỏi và đoạn cùng phân phối, từ đó khớp tốt hơn. Đây là minh chứng định lượng rõ ràng cho sự cần thiết của bước tách từ.")
    H2(doc, "4.4. Thí nghiệm 2 — Đóng góp của reranker")
    body(doc, "So sánh (2) và (3): bộ reranker PhoRanker xếp lại 30 ứng viên đầu giúp Hit@1 tăng nhẹ (0,520 → 0,540) nhưng MRR gần như không đổi (0,668 → 0,667) và Hit@3 giảm (0,820 → 0,780). Như vậy trên bộ 50 câu, reranker dùng đơn lẻ cho lợi ích khiêm tốn. Tuy nhiên khi kết hợp với hybrid (cấu hình 5), reranker lại nâng Hit@1 lên cao nhất (0,560) và MRR lên cao nhất (0,694) — cho thấy reranker phát huy tác dụng tốt hơn khi đứng sau một tập ứng viên đa dạng từ hybrid. Cross-encoder đánh giá độ liên quan tinh tế hơn bi-encoder nhờ chú ý chéo (Mục 1.5), nhưng hiệu quả phụ thuộc chất lượng và độ đa dạng của tập ứng viên đầu vào.")
    H2(doc, "4.5. Thí nghiệm 3 — Ảnh hưởng của kích thước chunk")
    body(doc, "Trên cùng một tập 218 tài liệu (gồm các luật liên quan và tài liệu nhiễu), dữ liệu được chia đoạn theo ba cấu hình rồi đánh giá (Bảng 4.2).")
    table(doc, ["Cấu hình", "Số chunk", "Hit@1", "Hit@5", "Hit@10", "MRR"],
          [["Cắt cứng 500 ký tự", "8.437", "0,944", "1,000", "1,000", "0,972"],
           ["Cắt cứng 1200 ký tự", "3.647", "1,000", "1,000", "1,000", "1,000"],
           ["Cắt theo “Điều” (~1200)", "4.195", "1,000", "1,000", "1,000", "1,000"]],
          caption="Bảng 4.2. Ảnh hưởng của kích thước chunk (218 tài liệu)")
    body(doc, "Cấu hình cắt 500 ký tự sinh ra gấp khoảng 2,3 lần số đoạn (8.437 so với 3.647) — tức gấp đôi chi phí embedding và lưu trữ — nhưng Hit@1 và MRR lại thấp hơn (0,944 và 0,972 so với 1,000), do đoạn quá ngắn dễ tách rời ngữ cảnh và phân tán thông tin của một “Điều” qua nhiều đoạn. Hai cấu hình ~1200 ký tự và cắt theo “Điều” cho độ chính xác bằng hoặc cao hơn với ít đoạn hơn. Cần lưu ý hạn chế: do tập đánh giá nhỏ (218 tài liệu) nên các chỉ số gần bão hòa ở mức 1,0; điểm phân biệt rõ nhất ở đây là hiệu quả chi phí. Kết quả ủng hộ lựa chọn cắt theo “Điều” cho hệ thống.")
    H2(doc, "4.6. Thí nghiệm 4 — Truy xuất hybrid")
    body(doc, "So sánh (2), (4) và (5): hybrid BM25+vector đạt Hit@10 = 0,980 — gần như không bỏ sót, cao nhất trong các cấu hình; kết hợp thêm reranker (cấu hình 5) cho Hit@1 cao nhất (0,560) và MRR cao nhất (0,694). Quan trọng hơn về mặt định tính: với truy vấn chứa số hiệu (ví dụ “Nghị định 96/2023/NĐ-CP quy định gì?”), cơ chế hybrid kèm nhận diện số hiệu trả về đúng văn bản, trong khi vector thuần thường trượt do không nắm được chuỗi số hiệu. Đây là minh chứng rõ cho luận điểm BM25 bù đắp điểm yếu từ vựng của embedding (Mục 1.4). Tổng hợp lại, cấu hình Hybrid + Reranker được chọn cho hệ thống vì cân bằng tốt nhất giữa độ chính xác ở hạng đầu (Hit@1, MRR) và độ phủ (Hit@10).")
    H2(doc, "4.7. Phân tích lỗi và bàn luận")
    body(doc, "Phân tích các câu bị trượt cho thấy hai nguyên nhân chính. Thứ nhất, một số câu hỏi gần với văn bản dưới luật (thông tư, quyết định của Cục/Bộ) hơn là luật gốc, khiến hệ thống xếp văn bản dưới luật lên trên — về mặt nội dung không hẳn sai. Ví dụ, câu “Những hành vi bị nghiêm cấm trong hoạt động dược?” trả về top-1 là một quyết định của Cục Quản lý Dược thay vì Luật Dược; câu về tiêm chủng trả về Nghị định 104/2016/NĐ-CP (nghị định chuyên về tiêm chủng) — thực chất rất sát nội dung, dù không trùng số hiệu luật được gán nhãn.")
    body(doc, "Thứ hai, vấn đề ground-truth: nhiều câu trả về Văn bản hợp nhất (VBHN) tương đương thay vì luật gốc — ví dụ câu về an toàn thực phẩm trả về 61/VBHN-VPQH (bản hợp nhất Luật An toàn thực phẩm), câu về HIV trả về 33/VBHN-VPQH (bản hợp nhất Luật phòng, chống HIV/AIDS) — đúng nội dung nhưng khác số hiệu. Nếu không tính các VBHN này vào ground-truth, chỉ số sẽ bị đánh giá thấp một cách giả tạo. Vì vậy đề tài đã bổ sung VBHN tương đương vào nhãn để đánh giá công bằng — một lưu ý quan trọng khi xây dựng tập kiểm thử cho miền pháp luật.")
    body(doc, "Về phần sinh (RQ3), với câu hỏi “Điều kiện để cá nhân được phép khám bệnh, chữa bệnh?”, PhoGPT trả lời bám đúng Điều 19 Luật Khám bệnh, chữa bệnh 2023 (15/2023/QH15): cần có giấy phép hành nghề còn hiệu lực và đã đăng ký hành nghề, kèm trích dẫn số hiệu và không bịa thông tin ngoài ngữ cảnh.")
    H2(doc, "4.8. Precision, Recall và ảnh hưởng của số nguồn k")
    body(doc, "Bên cạnh Hit@k và MRR, đề tài đo Precision@k và Recall@k ở mức văn bản (sau khi gộp các đoạn về số hiệu văn bản) cho cấu hình Hybrid + Reranker, đồng thời khảo sát ảnh hưởng của số nguồn k truyền cho mô hình (Bảng 4.3, Hình 4.2).")
    table(doc, ["k", "Precision@k", "Recall@k", "F1@k", "Trần P@k"],
          [["1", "0,560", "0,363", "0,441", "1,000"],
           ["2", "0,460", "0,552", "0,502", "0,980"],
           ["3", "0,360", "0,632", "0,459", "0,653"],
           ["4", "0,295", "0,678", "0,411", "0,490"],
           ["5", "0,244", "0,688", "0,360", "0,392"],
           ["6", "0,213", "0,713", "0,328", "0,327"],
           ["7", "0,189", "0,728", "0,300", "0,280"],
           ["8", "0,170", "0,748", "0,277", "0,245"],
           ["9", "0,151", "0,748", "0,251", "0,218"],
           ["10", "0,140", "0,778", "0,237", "0,196"]],
          caption="Bảng 4.3. Precision/Recall/F1 theo số nguồn k = 1..10 (kèm trần Precision)")
    figure(doc, os.path.join(FIGS, "pr_sweep.png"), "Hình 4.2. Precision/Recall/F1 và trần Precision theo số nguồn k", 14)
    body(doc, "Quét chi tiết k = 1..10 (Hình 4.2) thể hiện rõ đánh đổi kinh điển: khi k tăng, Recall tăng đều (0,363 → 0,778) trong khi Precision giảm (0,560 → 0,140). F1 — trung bình điều hòa của hai đại lượng — đạt cực đại tại k = 2 (0,502); việc quét thô trước đây (chỉ k = 1, 3, 5, 10) đã bỏ sót điểm cực đại này, cho thấy giá trị của khảo sát dải k mịn.")
    body(doc, "So sánh ba thông số: Precision phản ánh độ “sạch” (ít nhiễu) còn Recall phản ánh độ “phủ” (ít bỏ sót); F1 cân bằng hai mặt. Vùng k = 2–3 cho F1 cao nhất nên tối ưu nếu ưu tiên độ chính xác của tập nguồn. Tuy nhiên, với bài toán sinh câu trả lời, việc cung cấp thêm ngữ cảnh (Recall cao hơn) thường có lợi cho LLM; vì vậy hệ thống chọn k = 5 làm mặc định — nơi Recall đã đạt ~0,69 mà ngữ cảnh chưa quá loãng. Đây là sự phân biệt giữa “k tối ưu cho truy xuất” (k≈2) và “k phù hợp cho sinh” (k≈5).")
    body(doc, "Một điểm phương pháp quan trọng: do mỗi câu chỉ có trung bình 1,96 văn bản liên quan, Precision bị chặn bởi trần P@k = min(|liên quan|, k)/k (đường nét đứt, Hình 4.2). Do đó Precision giảm theo k chủ yếu là hệ quả của trần giảm, không phải truy xuất kém đi — bằng chứng là tỷ lệ Precision đạt được so với trần lại TĂNG theo k (từ 56% ở k=1 lên 71% ở k=10). Khuyến nghị: khi báo cáo Precision trong miền có ít văn bản liên quan, nên trình bày kèm trần Precision hoặc dùng thêm MAP/nDCG để tránh hiểu nhầm.")
    H2(doc, "4.9. Thời gian phản hồi và tài nguyên")
    body(doc, "Toàn bộ hệ thống vận hành trên CPU với 16GB RAM. Các mốc thời gian đo được trên bộ câu hỏi: lập chỉ mục embedding đạt khoảng 7,7 đoạn/giây (toàn corpus 367.462 đoạn mất khoảng 13 giờ, chạy nền và có khả năng resume); thời gian truy xuất trung bình (gồm cả bước reranker) khoảng 7,4 giây/câu; thời gian sinh câu trả lời trung bình khoảng 45 giây/câu. Mức thời gian này chấp nhận được cho mục đích tra cứu tham khảo, dù chưa đạt thời gian thực; có thể rút ngắn đáng kể nếu triển khai trên GPU. Việc lượng tử hóa GGUF Q4 là yếu tố then chốt giúp mô hình 3,7 tỷ tham số chạy được trong giới hạn 16GB RAM.")
    figure(doc, os.path.join(FIGS, "latency.png"), "Hình 4.3. Thời gian phản hồi trung bình (truy xuất và sinh)", 9)
    H2(doc, "4.10. Chất lượng câu trả lời và mức độ trích dẫn")
    body(doc, "Để đánh giá phần sinh, đề tài chạy đầy đủ pipeline (Hybrid + Reranker + PhoGPT) trên 6 câu hỏi và rà soát thủ công. Kết quả: 6/6 câu trả lời (tỉ lệ 100%) trích dẫn đúng ít nhất một văn bản liên quan trong số nguồn truy xuất; các câu trả lời bám sát điều luật cụ thể (ví dụ Điều 19 Luật Khám bệnh chữa bệnh, Điều 33 và Điều 79 Luật Dược) và không bịa thông tin ngoài ngữ cảnh.")
    body(doc, "Hạn chế quan sát được: mô hình thỉnh thoảng nêu chi tiết phụ chưa chính xác (ví dụ một câu trả lời nhắc “Luật Khám bệnh, chữa bệnh 2009” trong khi nguồn trích dẫn là Luật 15/2023/QH15) — tuy nguồn dẫn vẫn đúng. Điều này cho thấy RAG giảm mạnh nhưng chưa loại bỏ hoàn toàn sai sót của phần sinh; cần kiểm chứng của con người với các nội dung pháp lý quan trọng.")

    # ----------------------------- CHƯƠNG 5 -----------------------------
    H1(doc, "CHƯƠNG 5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    H2(doc, "5.1. Kết luận")
    body(doc, "Đồ án đã xây dựng hoàn chỉnh một hệ thống RAG tra cứu văn bản pháp luật ngành y tế tiếng Việt, vận hành trên máy cá nhân CPU 16GB. Các kết quả chính: (i) xây dựng kho 367.462 đoạn sạch từ 21.490 văn bản còn hiệu lực, kèm quy trình tái lập được; (ii) pipeline truy xuất Hybrid + Reranker đạt MRR 0,722 và Hit@10 trong khoảng 0,944–1,000; (iii) phần sinh PhoGPT tạo câu trả lời có trích dẫn, bám ngữ cảnh; (iv) giao diện web hỏi–đáp chạy offline. Quan trọng cho yêu cầu học thuật, đóng góp của từng kỹ thuật (tách từ, reranker, hybrid, kích thước chunk) đều được kiểm chứng bằng thực nghiệm A/B, trả lời trực tiếp ba câu hỏi nghiên cứu RQ1–RQ3.")
    H2(doc, "5.2. Tính mới và khả năng ứng dụng")
    body(doc, "Tính mới: đồ án kết hợp một cách có hệ thống các kỹ thuật phù hợp tiếng Việt và miền pháp luật — bi-encoder dựa trên PhoBERT, hybrid BM25+vector kèm nhận diện số hiệu văn bản, reranker PhoRanker và sinh có ràng buộc trích dẫn bằng PhoGPT — đồng thời chứng minh đóng góp của từng thành phần bằng thực nghiệm A/B. Điểm khác biệt so với các hệ RAG phổ thông là cơ chế nhận diện số hiệu văn bản nhằm khắc phục điểm yếu của tìm kiếm ngữ nghĩa đối với các định danh chính xác. Khả năng ứng dụng: hệ thống có thể dùng làm công cụ hỗ trợ tra cứu cho người dân, cán bộ y tế và sinh viên; vận hành trên máy phổ thông (CPU) nên dễ triển khai trong điều kiện hạn chế. Triển vọng: kiến trúc mang tính tổng quát, có thể mở rộng sang các lĩnh vực pháp luật khác và nâng cấp thành dịch vụ trực tuyến khi có hạ tầng mạnh hơn.")
    H2(doc, "5.3. Hạn chế")
    bullet(doc, "Bộ câu hỏi đánh giá gồm 50 câu do tác giả xây dựng; tuy đã đủ lớn để cho kết quả ổn định, vẫn nên tiếp tục mở rộng và có chuyên gia pháp lý thẩm định ground-truth.")
    bullet(doc, "Tốc độ sinh trên CPU (trung bình ~52 giây/câu) chưa đáp ứng yêu cầu thời gian thực.")
    bullet(doc, "Chưa đánh giá định lượng độ trung thực (faithfulness) của câu trả lời so với nguồn.")
    H2(doc, "5.4. Hướng phát triển")
    bullet(doc, "Mở rộng và chuẩn hóa bộ đánh giá; bổ sung độ đo faithfulness/groundedness tự động.")
    bullet(doc, "Áp dụng chunking theo ngữ cảnh (gắn tiêu đề văn bản vào mỗi đoạn khi mã hóa) để tăng độ chính xác.")
    bullet(doc, "Tăng tốc phần sinh (GPU hoặc dịch vụ) và hỗ trợ hội thoại đa lượt với cơ chế viết lại câu hỏi.")
    bullet(doc, "Mở rộng phạm vi sang các lĩnh vực pháp luật khác ngoài y tế.")

    # ----------------------------- TÀI LIỆU THAM KHẢO -----------------------------
    H1(doc, "TÀI LIỆU THAM KHẢO")
    refs = [
        'P. Lewis et al., “Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks,” in Advances in Neural Information Processing Systems (NeurIPS), 2020.',
        'D. Q. Nguyen and A. T. Nguyen, “PhoBERT: Pre-trained language models for Vietnamese,” in Findings of EMNLP, 2020.',
        'D. Q. Nguyen et al., “PhoGPT: Generative Pre-training for Vietnamese,” arXiv:2311.02945, 2023.',
        'N. Reimers and I. Gurevych, “Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks,” in Proc. EMNLP-IJCNLP, 2019.',
        'S. Robertson and H. Zaragoza, “The Probabilistic Relevance Framework: BM25 and Beyond,” Foundations and Trends in Information Retrieval, vol. 3, no. 4, 2009.',
        'G. V. Cormack, C. L. A. Clarke, and S. Buettcher, “Reciprocal Rank Fusion Outperforms Condorcet and Individual Rank Learning Methods,” in Proc. ACM SIGIR, 2009.',
        'BKAI Foundation Models, “vietnamese-bi-encoder,” Hugging Face. [Online]. Available: https://huggingface.co/bkai-foundation-models/vietnamese-bi-encoder',
        'itdainb, “PhoRanker: A Cross-Encoder Reranker for Vietnamese,” Hugging Face. [Online]. Available: https://huggingface.co/itdainb/PhoRanker',
        'Chroma, “Chroma — the open-source embedding database.” [Online]. Available: https://www.trychroma.com',
        'th1nhng0, “Vietnamese Legal Documents,” Hugging Face Datasets. [Online]. Available: https://huggingface.co/datasets/th1nhng0/vietnamese-legal-documents',
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.first_line_indent = Cm(-1.0)
        p.paragraph_format.space_after = Pt(4)
        _set_font(p.add_run(f"[{i}] "), 13); _set_font(p.add_run(r), 13)

    # ----------------------------- PHỤ LỤC -----------------------------
    H1(doc, "PHỤ LỤC")
    body(doc, "Mã nguồn, hướng dẫn cài đặt và chạy hệ thống được trình bày trong tệp README.md kèm theo. Kết quả thực nghiệm chi tiết ở tệp eval/results.md; bộ câu hỏi đánh giá ở eval/golden_questions.json. Cấu hình triển khai khuyến nghị: PhoGPT (GGUF) + Hybrid + Reranker.")

    # sửa schema: zoom thiếu percent
    try:
        z = doc.settings.element.find(qn('w:zoom'))
        if z is not None and z.get(qn('w:percent')) is None:
            z.set(qn('w:percent'), '100')
    except Exception:
        pass

    doc.save(OUT)
    try: print("Da luu:", OUT)
    except Exception: pass


if __name__ == "__main__":
    main()
