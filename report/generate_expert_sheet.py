# -*- coding: utf-8 -*-
"""Phiếu thẩm định ground-truth bộ câu hỏi đánh giá (cho chuyên gia/GVHD ký xác nhận).
Chạy: python report/generate_expert_sheet.py -> report/PhieuThamDinh_BoCauHoi.docx"""
import os, sys, json
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "PhieuThamDinh_BoCauHoi.docx")
FONT = "Times New Roman"
golden = json.load(open(os.path.join(os.path.dirname(HERE), "eval", "golden_questions.json"), encoding="utf-8"))


def sf(run, size=12, bold=False, italic=False):
    run.font.name = FONT; run.font.size = Pt(size); run.bold = bold; run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), FONT)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); tcPr.append(sh)


def main():
    doc = Document()
    n = doc.styles['Normal']; n.font.name = FONT; n.font.size = Pt(12)
    n.element.rPr.rFonts.set(qn('w:cs'), FONT)
    s = doc.sections[0]
    # khổ ngang cho bảng rộng
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width = Cm(29.7); s.page_height = Cm(21.0)
    s.top_margin = Cm(2); s.bottom_margin = Cm(2); s.left_margin = Cm(2.5); s.right_margin = Cm(2)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(t.add_run("PHIẾU THẨM ĐỊNH BỘ CÂU HỎI ĐÁNH GIÁ VÀ NHÃN VĂN BẢN LIÊN QUAN"), 14, bold=True)
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(st.add_run("Đề tài: Hệ thống RAG hỗ trợ tra cứu văn bản pháp luật ngành y tế Việt Nam — SV: Nguyễn Minh Hiếu"), 11, italic=True)
    p = doc.add_paragraph()
    sf(p.add_run("Kính đề nghị Thầy/Cô (chuyên gia) rà soát từng câu hỏi và nhãn “văn bản liên quan” (ground-truth) dưới đây; đánh dấu Đồng ý (✓) hoặc ghi đề xuất sửa vào cột Ghi chú. Loại câu: [T] soạn tay, [S] sinh từ tiêu đề Điều, [N] diễn đạt tự nhiên."), 11)

    tb = doc.add_table(rows=1, cols=5); tb.style = 'Table Grid'
    heads = ["#", "Câu hỏi", "Văn bản liên quan (nhãn)", "Loại", "Đồng ý / Ghi chú"]
    widths = [Cm(1.0), Cm(11.5), Cm(6.5), Cm(1.4), Cm(4.6)]
    for i, h in enumerate(heads):
        c = tb.rows[0].cells[i]; c.width = widths[i]; shade(c, "D9E2F3")
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sf(c.paragraphs[0].add_run(h), 11, bold=True)
    for i, it in enumerate(golden, 1):
        typ = "N" if it.get("type") == "natural" else ("T" if i <= 50 else "S")
        cells = tb.add_row().cells
        vals = [str(i), it["q"], ", ".join(it["relevant"]), typ, ""]
        for j, v in enumerate(vals):
            cells[j].width = widths[j]
            if j in (0, 3): cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            sf(cells[j].paragraphs[0].add_run(v), 10.5)
    for row in tb.rows:
        for j, c in enumerate(row.cells): c.width = widths[j]

    doc.add_paragraph()
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sf(p2.add_run("Xác nhận của chuyên gia/GVHD:  ........................................  Ngày ....../....../2026"), 12)

    try:
        z = doc.settings.element.find(qn('w:zoom'))
        if z is not None and z.get(qn('w:percent')) is None: z.set(qn('w:percent'), '100')
    except Exception: pass
    doc.save(OUT)
    print("Da luu:", OUT, f"({len(golden)} câu)")


if __name__ == "__main__":
    main()
